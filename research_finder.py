#!/usr/bin/env python3
"""
Advanced Research System (ARS) - Ultimate Literature Review & Analysis Tool
===========================================================================

A comprehensive research platform that combines:
- Literature review management with YAML metadata
- Multi-format document processing (PDF, DOCX, TXT, MD, HTML, etc.)
- Advanced full-text and metadata search
- Web research and crawling capabilities
- AI-powered summarization and analysis
- Automated report generation
- Multiple export formats
- Both GUI and CLI interfaces

Version: 2.0.0
Author: Advanced Research System Team
"""

import os
import sys
import re
import csv
import json
import yaml
import time
import math
import hashlib
import logging
import argparse
import threading
import subprocess
import webbrowser
import urllib.parse
import urllib.request
from pathlib import Path
from datetime import datetime
from collections import defaultdict, Counter, OrderedDict
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Dict, Tuple, Optional, Any, Union
import string
import traceback
import shutil
import textwrap
import sqlite3
import pickle

# Document processing imports
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    import docx2txt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from bs4 import BeautifulSoup, Comment
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False

# GUI imports
try:
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox, scrolledtext
    from tkinter.font import Font
    import tkinter.font as tkfont
    GUI_AVAILABLE = True
except ImportError:
    GUI_AVAILABLE = False

# Web framework imports
try:
    from flask import Flask, request, render_template, send_file, jsonify
    FLASK_AVAILABLE = True
except ImportError:
    FLASK_AVAILABLE = False

# AI/ML imports
try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False

# Initialize logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('ARS')

# Global configuration
DEFAULT_CONFIG = {
    # System settings
    "version": "2.0.0",
    "database_path": "research_database.db",
    "output_dir": "research_output",
    "temp_dir": "temp",
    "cache_dir": "cache",
    
    # Document processing
    "supported_formats": [".pdf", ".docx", ".txt", ".md", ".html", ".htm", ".rtf"],
    "extract_metadata": True,
    "extract_citations": True,
    "extract_figures": False,
    "ocr_enabled": False,
    
    # Search settings
    "search_method": "advanced",  # basic, advanced, semantic
    "case_sensitive": False,
    "whole_words": False,
    "fuzzy_matching": True,
    "relevance_threshold": 0.5,
    
    # Web research
    "enable_web_search": True,
    "search_engine": "duckduckgo",
    "max_web_results": 30,
    "crawl_depth": 2,
    "respect_robots_txt": True,
    
    # AI/Summarization
    "summarization_enabled": True,
    "summary_method": "hybrid",  # textrank, ollama, openai, hybrid
    "summary_model": "gemma2:9b",
    "ollama_url": "http://localhost:11434/api/generate",
    "openai_api_key": "",
    
    # Analysis settings
    "enable_statistics": True,
    "enable_clustering": True,
    "enable_topic_modeling": False,
    "citation_analysis": True,
    
    # Export settings
    "export_formats": ["xlsx", "csv", "json", "markdown", "html", "bibtex"],
    "include_visualizations": True,
    "generate_bibliography": True,
    
    # Performance
    "max_threads": 8,
    "batch_size": 10,
    "cache_enabled": True,
    "memory_limit_mb": 2048,
}

# Database schema
DATABASE_SCHEMA = """
CREATE TABLE IF NOT EXISTS documents (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    filename TEXT NOT NULL,
    filepath TEXT UNIQUE NOT NULL,
    filetype TEXT,
    filesize INTEGER,
    date_added TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    date_modified TIMESTAMP,
    content_hash TEXT,
    full_text TEXT,
    metadata TEXT,
    processed BOOLEAN DEFAULT FALSE,
    error_msg TEXT
);

CREATE TABLE IF NOT EXISTS metadata (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    document_id INTEGER,
    key TEXT,
    value TEXT,
    FOREIGN KEY (document_id) REFERENCES documents(id)
);

CREATE TABLE IF NOT EXISTS paragraphs (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    document_id INTEGER,
    page_num INTEGER,
    position INTEGER,
    content TEXT,
    FOREIGN KEY (document_id) REFERENCES documents(id)
);

CREATE TABLE IF NOT EXISTS search_results (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    search_query TEXT,
    document_id INTEGER,
    paragraph_id INTEGER,
    relevance_score REAL,
    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (document_id) REFERENCES documents(id),
    FOREIGN KEY (paragraph_id) REFERENCES paragraphs(id)
);

CREATE TABLE IF NOT EXISTS citations (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    document_id INTEGER,
    citation_text TEXT,
    citation_type TEXT,
    year INTEGER,
    authors TEXT,
    FOREIGN KEY (document_id) REFERENCES documents(id)
);

CREATE TABLE IF NOT EXISTS summaries (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    document_id INTEGER,
    summary_type TEXT,
    summary_text TEXT,
    method_used TEXT,
    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (document_id) REFERENCES documents(id)
);

CREATE TABLE IF NOT EXISTS web_sources (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    url TEXT UNIQUE,
    domain TEXT,
    title TEXT,
    content TEXT,
    crawl_depth INTEGER,
    parent_url TEXT,
    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

CREATE INDEX IF NOT EXISTS idx_metadata_key ON metadata(key);
CREATE INDEX IF NOT EXISTS idx_paragraphs_doc ON paragraphs(document_id);
CREATE INDEX IF NOT EXISTS idx_search_query ON search_results(search_query);
CREATE INDEX IF NOT EXISTS idx_citations_doc ON citations(document_id);
"""

class DatabaseManager:
    """Manages the research database"""
    
    def __init__(self, db_path: str):
        self.db_path = db_path
        self.init_database()
        self.migrate_database()
        
    def init_database(self):
        """Initialize the database with schema"""
        with sqlite3.connect(self.db_path) as conn:
            conn.executescript(DATABASE_SCHEMA)
            conn.commit()
    
    def migrate_database(self):
        """Migrate database schema to handle updates"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Check if date_modified column exists
            cursor.execute("PRAGMA table_info(documents)")
            columns = [col[1] for col in cursor.fetchall()]
            
            if 'date_modified' not in columns:
                logger.info("Migrating database: adding date_modified column")
                try:
                    cursor.execute("ALTER TABLE documents ADD COLUMN date_modified TIMESTAMP")
                    conn.commit()
                except sqlite3.OperationalError:
                    pass  # Column might already exist
            
            # Check other potentially missing columns
            if 'content_hash' not in columns:
                try:
                    cursor.execute("ALTER TABLE documents ADD COLUMN content_hash TEXT")
                    conn.commit()
                except sqlite3.OperationalError:
                    pass
            
            if 'processed' not in columns:
                try:
                    cursor.execute("ALTER TABLE documents ADD COLUMN processed BOOLEAN DEFAULT FALSE")
                    conn.commit()
                except sqlite3.OperationalError:
                    pass
            
            if 'error_msg' not in columns:
                try:
                    cursor.execute("ALTER TABLE documents ADD COLUMN error_msg TEXT")
                    conn.commit()
                except sqlite3.OperationalError:
                    pass
    
    def add_document(self, filepath: str, metadata: dict = None) -> int:
        """Add a document to the database"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Check if document already exists
            cursor.execute("SELECT id FROM documents WHERE filepath = ?", (filepath,))
            existing = cursor.fetchone()
            if existing:
                return existing[0]
            
            # Get file stats
            try:
                file_stat = os.stat(filepath)
                filesize = file_stat.st_size
                date_modified = datetime.fromtimestamp(file_stat.st_mtime)
            except:
                filesize = 0
                date_modified = datetime.now()
            
            # Insert document
            cursor.execute("""
                INSERT INTO documents (filename, filepath, filetype, filesize, date_modified)
                VALUES (?, ?, ?, ?, ?)
            """, (
                os.path.basename(filepath),
                filepath,
                os.path.splitext(filepath)[1].lower(),
                filesize,
                date_modified
            ))
            
            doc_id = cursor.lastrowid
            
            # Add metadata
            if metadata:
                for key, value in metadata.items():
                    cursor.execute("""
                        INSERT INTO metadata (document_id, key, value)
                        VALUES (?, ?, ?)
                    """, (doc_id, key, str(value)))
            
            conn.commit()
            return doc_id
    
    def update_document_content(self, doc_id: int, full_text: str, content_hash: str):
        """Update document content"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE documents 
                SET full_text = ?, content_hash = ?, processed = TRUE
                WHERE id = ?
            """, (full_text, content_hash, doc_id))
            conn.commit()
    
    def add_paragraphs(self, doc_id: int, paragraphs: List[Tuple[int, int, str]]):
        """Add paragraphs for a document"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.executemany("""
                INSERT INTO paragraphs (document_id, page_num, position, content)
                VALUES (?, ?, ?, ?)
            """, [(doc_id, p[0], p[1], p[2]) for p in paragraphs])
            conn.commit()
    
    def search_documents(self, query: str, search_type: str = "full_text") -> List[Dict]:
        """Search documents in database"""
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            if search_type == "full_text":
                cursor.execute("""
                    SELECT d.*, COUNT(p.id) as match_count
                    FROM documents d
                    LEFT JOIN paragraphs p ON d.id = p.document_id
                    WHERE d.full_text LIKE ? OR p.content LIKE ?
                    GROUP BY d.id
                    ORDER BY match_count DESC
                """, (f"%{query}%", f"%{query}%"))
            elif search_type == "metadata":
                cursor.execute("""
                    SELECT DISTINCT d.*
                    FROM documents d
                    JOIN metadata m ON d.id = m.document_id
                    WHERE m.value LIKE ?
                """, (f"%{query}%",))
            
            return [dict(row) for row in cursor.fetchall()]
    
    def get_document_metadata(self, doc_id: int) -> Dict:
        """Get all metadata for a document"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT key, value FROM metadata WHERE document_id = ?
            """, (doc_id,))
            return {row[0]: row[1] for row in cursor.fetchall()}
    
    def add_search_result(self, query: str, doc_id: int, para_id: int, score: float):
        """Record a search result"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO search_results (search_query, document_id, paragraph_id, relevance_score)
                VALUES (?, ?, ?, ?)
            """, (query, doc_id, para_id, score))
            conn.commit()

class DocumentProcessor:
    """Advanced document processing engine"""
    
    def __init__(self, config: dict, db_manager: DatabaseManager):
        self.config = config
        self.db = db_manager
        self.processors = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.html': self._process_html,
            '.htm': self._process_html,
        }
    
    def process_document(self, filepath: str) -> Dict:
        """Process a document and extract all information"""
        logger.info(f"Processing document: {filepath}")
        
        # Get file extension
        ext = os.path.splitext(filepath)[1].lower()
        
        if ext not in self.processors:
            return {"error": f"Unsupported file type: {ext}"}
        
        try:
            # Add to database
            doc_id = self.db.add_document(filepath)
            
            # Process based on type
            result = self.processors[ext](filepath)
            
            # Calculate content hash
            content_hash = hashlib.sha256(result['full_text'].encode()).hexdigest()
            
            # Update database
            self.db.update_document_content(doc_id, result['full_text'], content_hash)
            
            # Add paragraphs
            if 'paragraphs' in result:
                self.db.add_paragraphs(doc_id, result['paragraphs'])
            
            # Extract and store metadata
            if result.get('metadata'):
                for key, value in result['metadata'].items():
                    self.db.add_metadata(doc_id, key, value)
            
            result['doc_id'] = doc_id
            return result
            
        except Exception as e:
            logger.error(f"Error processing {filepath}: {str(e)}")
            return {"error": str(e)}
    
    def _process_pdf(self, filepath: str) -> Dict:
        """Process PDF files with multiple extraction methods"""
        result = {
            'filepath': filepath,
            'filetype': 'pdf',
            'full_text': '',
            'paragraphs': [],
            'metadata': {},
            'pages': 0
        }
        
        # Try PyMuPDF first (most accurate)
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(filepath)
                result['pages'] = len(doc)
                
                # Extract metadata
                metadata = doc.metadata
                if metadata:
                    result['metadata'] = {
                        'title': metadata.get('title', ''),
                        'author': metadata.get('author', ''),
                        'subject': metadata.get('subject', ''),
                        'keywords': metadata.get('keywords', ''),
                        'creator': metadata.get('creator', ''),
                        'producer': metadata.get('producer', ''),
                        'creation_date': str(metadata.get('creationDate', '')),
                        'modification_date': str(metadata.get('modDate', ''))
                    }
                
                # Extract text and paragraphs
                full_text = []
                for page_num, page in enumerate(doc, 1):
                    page_text = page.get_text()
                    full_text.append(page_text)
                    
                    # Extract paragraphs with positions
                    paragraphs = self._extract_paragraphs(page_text)
                    for para_text, position in paragraphs:
                        result['paragraphs'].append((page_num, position, para_text))
                
                result['full_text'] = '\n\n'.join(full_text)
                doc.close()
                
                # Extract YAML metadata if present
                yaml_meta = self._extract_yaml_metadata(result['full_text'])
                if yaml_meta:
                    result['metadata'].update(yaml_meta)
                
                return result
                
            except Exception as e:
                logger.warning(f"PyMuPDF failed for {filepath}: {str(e)}")
        
        # Fallback to pdfplumber
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(filepath) as pdf:
                    result['pages'] = len(pdf.pages)
                    
                    # Extract metadata
                    if pdf.metadata:
                        result['metadata'] = {k: str(v) for k, v in pdf.metadata.items()}
                    
                    # Extract text
                    full_text = []
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text() or ""
                        full_text.append(page_text)
                        
                        # Extract paragraphs
                        paragraphs = self._extract_paragraphs(page_text)
                        for para_text, position in paragraphs:
                            result['paragraphs'].append((page_num, position, para_text))
                    
                    result['full_text'] = '\n\n'.join(full_text)
                    
                    # Extract YAML metadata
                    yaml_meta = self._extract_yaml_metadata(result['full_text'])
                    if yaml_meta:
                        result['metadata'].update(yaml_meta)
                    
                    return result
                    
            except Exception as e:
                logger.warning(f"pdfplumber failed for {filepath}: {str(e)}")
        
        # Final fallback to PyPDF2
        if PYPDF2_AVAILABLE:
            try:
                with open(filepath, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    result['pages'] = len(reader.pages)
                    
                    # Extract metadata
                    if reader.metadata:
                        result['metadata'] = {
                            'title': getattr(reader.metadata, 'title', ''),
                            'author': getattr(reader.metadata, 'author', ''),
                            'subject': getattr(reader.metadata, 'subject', ''),
                            'creator': getattr(reader.metadata, 'creator', ''),
                        }
                    
                    # Extract text
                    full_text = []
                    for page_num, page in enumerate(reader.pages, 1):
                        page_text = page.extract_text()
                        full_text.append(page_text)
                        
                        # Extract paragraphs
                        paragraphs = self._extract_paragraphs(page_text)
                        for para_text, position in paragraphs:
                            result['paragraphs'].append((page_num, position, para_text))
                    
                    result['full_text'] = '\n\n'.join(full_text)
                    
                    # Extract YAML metadata
                    yaml_meta = self._extract_yaml_metadata(result['full_text'])
                    if yaml_meta:
                        result['metadata'].update(yaml_meta)
                    
                    return result
                    
            except Exception as e:
                logger.error(f"All PDF processors failed for {filepath}: {str(e)}")
                raise
        
        raise Exception("No PDF processing library available")
    
    def _process_docx(self, filepath: str) -> Dict:
        """Process Word documents"""
        result = {
            'filepath': filepath,
            'filetype': 'docx',
            'full_text': '',
            'paragraphs': [],
            'metadata': {},
            'pages': 0
        }
        
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not installed")
        
        try:
            # Extract with python-docx for structure
            doc = DocxDocument(filepath)
            
            # Extract metadata from core properties
            core_props = doc.core_properties
            result['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or ''
            }
            
            # Extract paragraphs
            full_text = []
            position = 0
            for para_num, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if para_text:
                    full_text.append(para_text)
                    result['paragraphs'].append((1, position, para_text))
                    position += len(para_text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            full_text.append(cell_text)
            
            result['full_text'] = '\n\n'.join(full_text)
            
            # Extract YAML metadata
            yaml_meta = self._extract_yaml_metadata(result['full_text'])
            if yaml_meta:
                result['metadata'].update(yaml_meta)
            
            return result
            
        except Exception as e:
            # Fallback to docx2txt
            try:
                text = docx2txt.process(filepath)
                result['full_text'] = text
                
                # Extract paragraphs
                paragraphs = self._extract_paragraphs(text)
                for para_text, position in paragraphs:
                    result['paragraphs'].append((1, position, para_text))
                
                # Extract YAML metadata
                yaml_meta = self._extract_yaml_metadata(text)
                if yaml_meta:
                    result['metadata'] = yaml_meta
                
                return result
            except Exception as e2:
                logger.error(f"Error processing DOCX {filepath}: {str(e2)}")
                raise
    
    def _process_text(self, filepath: str) -> Dict:
        """Process plain text files"""
        result = {
            'filepath': filepath,
            'filetype': 'txt',
            'full_text': '',
            'paragraphs': [],
            'metadata': {},
            'pages': 1
        }
        
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            result['full_text'] = text
            
            # Extract paragraphs
            paragraphs = self._extract_paragraphs(text)
            for para_text, position in paragraphs:
                result['paragraphs'].append((1, position, para_text))
            
            # Extract YAML metadata
            yaml_meta = self._extract_yaml_metadata(text)
            if yaml_meta:
                result['metadata'] = yaml_meta
            
            # Try to extract citation-style metadata
            citation_meta = self._extract_citation_metadata(text)
            if citation_meta:
                result['metadata'].update(citation_meta)
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing text file {filepath}: {str(e)}")
            raise
    
    def _process_markdown(self, filepath: str) -> Dict:
        """Process Markdown files"""
        result = {
            'filepath': filepath,
            'filetype': 'md',
            'full_text': '',
            'paragraphs': [],
            'metadata': {},
            'pages': 1
        }
        
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                md_text = f.read()
            
            result['full_text'] = md_text
            
            # Extract YAML frontmatter
            yaml_meta = self._extract_yaml_metadata(md_text)
            if yaml_meta:
                result['metadata'] = yaml_meta
            
            # Convert to HTML for structure analysis
            if MARKDOWN_AVAILABLE:
                html = markdown.markdown(md_text, extensions=['meta', 'tables', 'fenced_code'])
                
                # Extract text from HTML
                if BS4_AVAILABLE:
                    soup = BeautifulSoup(html, 'html.parser')
                    text = soup.get_text()
                    
                    # Extract paragraphs from the text
                    paragraphs = self._extract_paragraphs(text)
                    for para_text, position in paragraphs:
                        result['paragraphs'].append((1, position, para_text))
            else:
                # Fallback: extract paragraphs from raw markdown
                paragraphs = self._extract_paragraphs(md_text)
                for para_text, position in paragraphs:
                    result['paragraphs'].append((1, position, para_text))
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing Markdown file {filepath}: {str(e)}")
            raise
    
    def _process_html(self, filepath: str) -> Dict:
        """Process HTML files"""
        result = {
            'filepath': filepath,
            'filetype': 'html',
            'full_text': '',
            'paragraphs': [],
            'metadata': {},
            'pages': 1
        }
        
        if not BS4_AVAILABLE:
            raise Exception("BeautifulSoup4 not installed")
        
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                html = f.read()
            
            soup = BeautifulSoup(html, 'html.parser')
            
            # Extract metadata from meta tags
            meta_tags = soup.find_all('meta')
            for tag in meta_tags:
                name = tag.get('name', tag.get('property', ''))
                content = tag.get('content', '')
                if name and content:
                    result['metadata'][name] = content
            
            # Extract title
            title_tag = soup.find('title')
            if title_tag:
                result['metadata']['title'] = title_tag.get_text(strip=True)
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            result['full_text'] = text
            
            # Extract paragraphs
            paragraphs = []
            for p in soup.find_all(['p', 'div', 'section', 'article']):
                para_text = p.get_text(strip=True)
                if len(para_text) > 50:  # Minimum paragraph length
                    paragraphs.append(para_text)
            
            # Add paragraphs with positions
            position = 0
            for para_text in paragraphs:
                result['paragraphs'].append((1, position, para_text))
                position += len(para_text)
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing HTML file {filepath}: {str(e)}")
            raise
    
    def _extract_paragraphs(self, text: str, min_length: int = 50) -> List[Tuple[str, int]]:
        """Extract paragraphs from text with positions"""
        paragraphs = []
        
        # Split by double newlines or paragraph patterns
        splits = re.split(r'\n\s*\n|\r\n\s*\r\n', text)
        
        current_pos = 0
        for para in splits:
            para = para.strip()
            if len(para) >= min_length:
                # Find position in original text
                pos = text.find(para, current_pos)
                if pos != -1:
                    paragraphs.append((para, pos))
                    current_pos = pos + len(para)
        
        # If no paragraphs found with double newlines, try single newlines
        if not paragraphs:
            lines = text.split('\n')
            current_para = []
            current_pos = 0
            
            for line in lines:
                line = line.strip()
                if line:
                    current_para.append(line)
                elif current_para:
                    para_text = ' '.join(current_para)
                    if len(para_text) >= min_length:
                        pos = text.find(para_text, current_pos)
                        if pos != -1:
                            paragraphs.append((para_text, pos))
                            current_pos = pos + len(para_text)
                    current_para = []
            
            # Don't forget the last paragraph
            if current_para:
                para_text = ' '.join(current_para)
                if len(para_text) >= min_length:
                    pos = text.find(para_text, current_pos)
                    if pos != -1:
                        paragraphs.append((para_text, pos))
        
        return paragraphs
    
    def _extract_yaml_metadata(self, text: str) -> Dict:
        """Extract YAML metadata from text"""
        # Look for YAML front matter
        yaml_pattern = re.compile(r'^---\s*\n(.*?)\n---\s*\n', re.DOTALL | re.MULTILINE)
        match = yaml_pattern.search(text)
        
        if match:
            yaml_text = match.group(1)
            try:
                metadata = yaml.safe_load(yaml_text)
                if isinstance(metadata, dict):
                    # Normalize keys to lowercase with underscores
                    normalized = {}
                    for key, value in metadata.items():
                        norm_key = key.lower().replace(' ', '_').replace('-', '_')
                        normalized[norm_key] = value
                    return normalized
            except yaml.YAMLError as e:
                logger.warning(f"Error parsing YAML metadata: {str(e)}")
        
        return {}
    
    def _extract_citation_metadata(self, text: str) -> Dict:
        """Extract citation-style metadata from text"""
        metadata = {}
        
        # Common citation patterns
        patterns = {
            'author': r'(?:Author|Authors?|By):\s*(.+?)(?:\n|$)',
            'title': r'(?:Title):\s*(.+?)(?:\n|$)',
            'year': r'(?:Year|Date):\s*(\d{4})',
            'journal': r'(?:Journal|Published in):\s*(.+?)(?:\n|$)',
            'doi': r'(?:DOI|doi):\s*(.+?)(?:\n|$)',
            'abstract': r'(?:Abstract):\s*(.+?)(?:\n\n|$)',
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                metadata[key] = match.group(1).strip()
        
        return metadata

class AdvancedSearchEngine:
    """Advanced search engine with multiple search methods"""
    
    def __init__(self, config: dict, db_manager: DatabaseManager):
        self.config = config
        self.db = db_manager
        self.search_methods = {
            'basic': self._basic_search,
            'advanced': self._advanced_search,
            'semantic': self._semantic_search
        }
    
    def search(self, query: str, method: str = None, filters: Dict = None) -> List[Dict]:
        """Perform search with specified method"""
        method = method or self.config.get('search_method', 'advanced')
        
        if method not in self.search_methods:
            method = 'advanced'
        
        logger.info(f"Performing {method} search for: {query}")
        
        # Execute search
        results = self.search_methods[method](query, filters)
        
        # Apply post-processing
        results = self._apply_relevance_ranking(results, query)
        
        # Filter by threshold
        threshold = self.config.get('relevance_threshold', 0.5)
        results = [r for r in results if r.get('relevance_score', 0) >= threshold]
        
        # Record search results in database
        for result in results:
            self.db.add_search_result(
                query, 
                result['doc_id'], 
                result.get('paragraph_id'),
                result['relevance_score']
            )
        
        return results
    
    def _basic_search(self, query: str, filters: Dict = None) -> List[Dict]:
        """Basic keyword search"""
        results = []
        
        # Search in database
        db_results = self.db.search_documents(query, 'full_text')
        
        for doc in db_results:
            result = {
                'doc_id': doc['id'],
                'filename': doc['filename'],
                'filepath': doc['filepath'],
                'match_count': doc.get('match_count', 0),
                'relevance_score': 0.5  # Basic score
            }
            results.append(result)
        
        return results
    
    def _advanced_search(self, query: str, filters: Dict = None) -> List[Dict]:
        """Advanced search with query parsing and ranking"""
        results = []
        
        # Parse query for special operators
        parsed_query = self._parse_advanced_query(query)
        
        # Build SQL query based on parsed query
        with sqlite3.connect(self.db.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Complex query with relevance scoring
            sql = """
                SELECT 
                    d.id, d.filename, d.filepath,
                    p.id as paragraph_id, p.content as paragraph_content,
                    p.page_num, p.position,
                    COUNT(DISTINCT p.id) as match_count,
                    SUM(
                        CASE 
                            WHEN p.content LIKE ? THEN 3
                            WHEN p.content LIKE ? THEN 2
                            WHEN d.full_text LIKE ? THEN 1
                            ELSE 0
                        END
                    ) as relevance_score
                FROM documents d
                LEFT JOIN paragraphs p ON d.id = p.document_id
                WHERE 1=1
            """
            
            params = []
            
            # Add search conditions
            for term in parsed_query['must_include']:
                sql += " AND (p.content LIKE ? OR d.full_text LIKE ?)"
                params.extend([f"%{term}%", f"%{term}%"])
            
            for term in parsed_query['must_exclude']:
                sql += " AND NOT (p.content LIKE ? OR d.full_text LIKE ?)"
                params.extend([f"%{term}%", f"%{term}%"])
            
            # Add relevance parameters
            params = [f"%{query}%", f"% {query} %", f"%{query}%"] + params
            
            sql += " GROUP BY d.id, p.id ORDER BY relevance_score DESC, match_count DESC"
            
            cursor.execute(sql, params)
            
            for row in cursor.fetchall():
                result = {
                    'doc_id': row['id'],
                    'filename': row['filename'],
                    'filepath': row['filepath'],
                    'paragraph_id': row['paragraph_id'],
                    'paragraph_content': row['paragraph_content'],
                    'page_num': row['page_num'],
                    'position': row['position'],
                    'match_count': row['match_count'],
                    'relevance_score': row['relevance_score'] / 10.0  # Normalize
                }
                results.append(result)
        
        return results
    
    def _semantic_search(self, query: str, filters: Dict = None) -> List[Dict]:
        """Semantic search using embeddings (placeholder for future implementation)"""
        # This would use sentence embeddings for semantic similarity
        # For now, fall back to advanced search
        return self._advanced_search(query, filters)
    
    def _parse_advanced_query(self, query: str) -> Dict:
        """Parse advanced query syntax"""
        parsed = {
            'must_include': [],
            'should_include': [],
            'must_exclude': [],
            'exact_phrases': []
        }
        
        # Extract exact phrases (in quotes)
        exact_pattern = r'"([^"]+)"'
        for match in re.finditer(exact_pattern, query):
            parsed['exact_phrases'].append(match.group(1))
            query = query.replace(match.group(0), '')
        
        # Extract must include (+term)
        must_pattern = r'\+(\S+)'
        for match in re.finditer(must_pattern, query):
            parsed['must_include'].append(match.group(1))
            query = query.replace(match.group(0), '')
        
        # Extract must exclude (-term)
        exclude_pattern = r'-(\S+)'
        for match in re.finditer(exclude_pattern, query):
            parsed['must_exclude'].append(match.group(1))
            query = query.replace(match.group(0), '')
        
        # Remaining terms are should include
        remaining_terms = query.strip().split()
        parsed['should_include'] = [t for t in remaining_terms if t]
        
        # If no special operators, treat all as must include
        if not any([parsed['must_include'], parsed['should_include'], 
                   parsed['exact_phrases'], parsed['must_exclude']]):
            parsed['must_include'] = query.strip().split()
        
        return parsed
    
    def _apply_relevance_ranking(self, results: List[Dict], query: str) -> List[Dict]:
        """Apply advanced relevance ranking to results"""
        query_terms = query.lower().split()
        
        for result in results:
            score = result.get('relevance_score', 0)
            
            # Boost for title matches
            if 'filename' in result:
                filename_lower = result['filename'].lower()
                for term in query_terms:
                    if term in filename_lower:
                        score += 0.3
            
            # Boost for metadata matches
            if 'metadata' in result:
                meta_text = ' '.join(str(v) for v in result['metadata'].values()).lower()
                for term in query_terms:
                    if term in meta_text:
                        score += 0.2
            
            # Boost for exact phrase matches
            if 'paragraph_content' in result:
                para_lower = result['paragraph_content'].lower()
                if query.lower() in para_lower:
                    score += 0.5
            
            result['relevance_score'] = min(score, 1.0)  # Cap at 1.0
        
        # Sort by relevance
        return sorted(results, key=lambda x: x['relevance_score'], reverse=True)

class WebResearchEngine:
    """Web research and crawling engine"""
    
    def __init__(self, config: dict, db_manager: DatabaseManager):
        self.config = config
        self.db = db_manager
        self.session = requests.Session() if REQUESTS_AVAILABLE else None
        self.visited_urls = set()
    
    def search_web(self, query: str, max_results: int = None) -> List[Dict]:
        """Search the web for information"""
        max_results = max_results or self.config.get('max_web_results', 30)
        results = []
        
        if not self.config.get('enable_web_search', True):
            return results
        
        logger.info(f"Searching web for: {query}")
        
        # Use DuckDuckGo HTML API
        try:
            search_url = f"https://html.duckduckgo.com/html/?q={urllib.parse.quote(query)}"
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(search_url, headers=headers, timeout=10)
            response.raise_for_status()
            
            if BS4_AVAILABLE:
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Extract search results
                for result_div in soup.find_all('div', class_='result'):
                    title_elem = result_div.find('h2', class_='result__title')
                    snippet_elem = result_div.find('a', class_='result__snippet')
                    
                    if title_elem and snippet_elem:
                        link = title_elem.find('a')
                        if link and link.get('href'):
                            url = link['href']
                            # Extract actual URL from DuckDuckGo redirect
                            if 'uddg=' in url:
                                url = urllib.parse.unquote(url.split('uddg=')[1].split('&')[0])
                            
                            results.append({
                                'url': url,
                                'title': title_elem.get_text(strip=True),
                                'snippet': snippet_elem.get_text(strip=True),
                                'source': 'duckduckgo'
                            })
                            
                            if len(results) >= max_results:
                                break
            
        except Exception as e:
            logger.error(f"Web search error: {str(e)}")
        
        return results
    
    def crawl_url(self, url: str, depth: int = 1) -> Dict:
        """Crawl a URL and extract content"""
        if url in self.visited_urls:
            return {'url': url, 'already_visited': True}
        
        self.visited_urls.add(url)
        
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (compatible; ResearchBot/1.0)'
            }
            
            response = requests.get(url, headers=headers, timeout=15)
            response.raise_for_status()
            
            # Parse content
            content_type = response.headers.get('content-type', '').lower()
            
            if 'text/html' in content_type and BS4_AVAILABLE:
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Extract title
                title = ''
                title_tag = soup.find('title')
                if title_tag:
                    title = title_tag.get_text(strip=True)
                
                # Remove scripts and styles
                for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
                    tag.decompose()
                
                # Extract main content
                main_content = None
                for selector in ['main', 'article', '#content', '.content', 'body']:
                    main_content = soup.find(selector)
                    if main_content:
                        break
                
                if not main_content:
                    main_content = soup.find('body')
                
                # Extract text
                text = main_content.get_text(separator='\n', strip=True) if main_content else ''
                
                # Extract links for further crawling
                links = []
                if depth > 0:
                    for link in soup.find_all('a', href=True):
                        abs_url = urllib.parse.urljoin(url, link['href'])
                        if self._is_valid_url(abs_url) and abs_url not in self.visited_urls:
                            links.append(abs_url)
                
                # Store in database
                with sqlite3.connect(self.db.db_path) as conn:
                    cursor = conn.cursor()
                    cursor.execute("""
                        INSERT OR REPLACE INTO web_sources 
                        (url, domain, title, content, crawl_depth)
                        VALUES (?, ?, ?, ?, ?)
                    """, (
                        url,
                        urllib.parse.urlparse(url).netloc,
                        title,
                        text,
                        depth
                    ))
                    conn.commit()
                
                return {
                    'url': url,
                    'title': title,
                    'content': text,
                    'links': links[:10],  # Limit links
                    'success': True
                }
            
        except Exception as e:
            logger.error(f"Error crawling {url}: {str(e)}")
            return {'url': url, 'error': str(e), 'success': False}
    
    def _is_valid_url(self, url: str) -> bool:
        """Check if URL is valid for crawling"""
        try:
            parsed = urllib.parse.urlparse(url)
            
            # Check scheme
            if parsed.scheme not in ['http', 'https']:
                return False
            
            # Check for common file extensions to skip
            skip_extensions = {'.pdf', '.jpg', '.jpeg', '.png', '.gif', '.zip', '.exe'}
            if any(url.lower().endswith(ext) for ext in skip_extensions):
                return False
            
            return True
            
        except:
            return False

class SummarizationEngine:
    """Advanced summarization engine with multiple methods"""
    
    def __init__(self, config: dict):
        self.config = config
        self.methods = {
            'textrank': self._summarize_textrank,
            'frequency': self._summarize_frequency,
            'ollama': self._summarize_ollama,
            'openai': self._summarize_openai,
            'hybrid': self._summarize_hybrid
        }
    
    def summarize(self, text: str, method: str = None, length: str = 'medium') -> str:
        """Summarize text using specified method"""
        if not text or len(text) < 100:
            return text
        
        method = method or self.config.get('summary_method', 'hybrid')
        
        if method not in self.methods:
            method = 'frequency'
        
        try:
            return self.methods[method](text, length)
        except Exception as e:
            logger.error(f"Summarization error with {method}: {str(e)}")
            # Fallback to frequency-based
            return self._summarize_frequency(text, length)
    
    def _summarize_textrank(self, text: str, length: str) -> str:
        """TextRank summarization"""
        if not NLTK_AVAILABLE:
            return self._summarize_frequency(text, length)
        
        try:
            # Tokenize into sentences
            sentences = sent_tokenize(text)
            
            if len(sentences) < 3:
                return text
            
            # Calculate sentence scores using TextRank algorithm
            # (Simplified version)
            word_freq = Counter()
            for sentence in sentences:
                words = word_tokenize(sentence.lower())
                words = [w for w in words if w.isalnum() and len(w) > 3]
                word_freq.update(words)
            
            # Score sentences
            sentence_scores = {}
            for sentence in sentences:
                words = word_tokenize(sentence.lower())
                words = [w for w in words if w.isalnum() and len(w) > 3]
                
                if len(words) > 0:
                    score = sum(word_freq[w] for w in words) / len(words)
                    sentence_scores[sentence] = score
            
            # Select top sentences
            num_sentences = {
                'short': max(1, len(sentences) // 5),
                'medium': max(2, len(sentences) // 3),
                'long': max(3, len(sentences) // 2)
            }.get(length, len(sentences) // 3)
            
            top_sentences = sorted(sentence_scores.items(), 
                                 key=lambda x: x[1], 
                                 reverse=True)[:num_sentences]
            
            # Reorder by original position
            summary_sentences = []
            for sentence in sentences:
                if any(sentence == s[0] for s in top_sentences):
                    summary_sentences.append(sentence)
            
            return ' '.join(summary_sentences)
            
        except Exception as e:
            logger.error(f"TextRank error: {str(e)}")
            return self._summarize_frequency(text, length)
    
    def _summarize_frequency(self, text: str, length: str) -> str:
        """Simple frequency-based summarization"""
        # Split into sentences
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
        
        if len(sentences) < 3:
            return text
        
        # Calculate word frequencies
        words = re.findall(r'\b\w+\b', text.lower())
        word_freq = Counter(words)
        
        # Remove common words
        common_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 
                       'to', 'for', 'of', 'with', 'by', 'from', 'is', 'was', 'are', 'were'}
        for word in common_words:
            word_freq.pop(word, None)
        
        # Score sentences
        sentence_scores = {}
        for sentence in sentences:
            words = re.findall(r'\b\w+\b', sentence.lower())
            if words:
                score = sum(word_freq.get(w, 0) for w in words) / len(words)
                sentence_scores[sentence] = score
        
        # Select top sentences
        num_sentences = {
            'short': max(1, len(sentences) // 5),
            'medium': max(2, len(sentences) // 3),
            'long': max(3, len(sentences) // 2)
        }.get(length, len(sentences) // 3)
        
        top_sentences = sorted(sentence_scores.items(), 
                             key=lambda x: x[1], 
                             reverse=True)[:num_sentences]
        
        # Return sentences in original order
        summary = []
        for sentence in sentences:
            if any(sentence == s[0] for s in top_sentences):
                summary.append(sentence)
        
        return '. '.join(summary) + '.'
    
    def _summarize_ollama(self, text: str, length: str) -> str:
        """Summarize using Ollama"""
        if not REQUESTS_AVAILABLE:
            return self._summarize_frequency(text, length)
        
        try:
            prompt = f"""Summarize the following text in {length} detail:

{text[:8000]}  # Limit context

Summary:"""
            
            response = requests.post(
                self.config.get('ollama_url'),
                json={
                    'model': self.config.get('summary_model', 'mistral'),
                    'prompt': prompt,
                    'stream': False
                },
                timeout=30
            )
            
            if response.status_code == 200:
                return response.json().get('response', '').strip()
            else:
                raise Exception(f"Ollama API error: {response.status_code}")
                
        except Exception as e:
            logger.error(f"Ollama summarization error: {str(e)}")
            return self._summarize_frequency(text, length)
    
    def _summarize_openai(self, text: str, length: str) -> str:
        """Summarize using OpenAI"""
        api_key = self.config.get('openai_api_key')
        if not api_key or not REQUESTS_AVAILABLE:
            return self._summarize_frequency(text, length)
        
        try:
            headers = {
                'Authorization': f'Bearer {api_key}',
                'Content-Type': 'application/json'
            }
            
            messages = [
                {"role": "system", "content": f"You are a research assistant. Provide a {length} summary."},
                {"role": "user", "content": f"Summarize this text:\n\n{text[:8000]}"}
            ]
            
            response = requests.post(
                'https://api.openai.com/v1/chat/completions',
                headers=headers,
                json={
                    'model': 'gpt-3.5-turbo',
                    'messages': messages,
                    'temperature': 0.3,
                    'max_tokens': 500 if length == 'short' else 1000
                },
                timeout=30
            )
            
            if response.status_code == 200:
                return response.json()['choices'][0]['message']['content'].strip()
            else:
                raise Exception(f"OpenAI API error: {response.status_code}")
                
        except Exception as e:
            logger.error(f"OpenAI summarization error: {str(e)}")
            return self._summarize_frequency(text, length)
    
    def _summarize_hybrid(self, text: str, length: str) -> str:
        """Hybrid summarization combining multiple methods"""
        summaries = []
        
        # Try AI methods first
        if self.config.get('ollama_url'):
            summary = self._summarize_ollama(text, length)
            if summary and len(summary) > 50:
                summaries.append(summary)
        
        if self.config.get('openai_api_key'):
            summary = self._summarize_openai(text, length)
            if summary and len(summary) > 50:
                summaries.append(summary)
        
        # Always include extractive summary
        extractive = self._summarize_textrank(text, length)
        if extractive and len(extractive) > 50:
            summaries.append(extractive)
        
        if not summaries:
            return self._summarize_frequency(text, length)
        
        # Combine summaries intelligently
        if len(summaries) == 1:
            return summaries[0]
        
        # If multiple summaries, combine key points
        combined = "Summary:\n\n"
        for i, summary in enumerate(summaries, 1):
            if len(summaries) > 1:
                combined += f"[Method {i}] "
            combined += summary + "\n\n"
        
        return combined.strip()

class ReportGenerator:
    """Advanced report generation engine"""
    
    def __init__(self, config: dict):
        self.config = config
        self.formats = {
            'markdown': self._generate_markdown,
            'html': self._generate_html,
            'pdf': self._generate_pdf,
            'docx': self._generate_docx,
            'latex': self._generate_latex
        }
    
    def generate_report(self, research_data: Dict, format: str = 'markdown') -> str:
        """Generate comprehensive research report"""
        if format not in self.formats:
            format = 'markdown'
        
        return self.formats[format](research_data)
    
    def _generate_markdown(self, data: Dict) -> str:
        """Generate Markdown report"""
        report = []
        
        # Header
        report.append(f"# {data.get('title', 'Research Report')}")
        report.append(f"\n*Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n")
        
        # Executive Summary
        if data.get('summary'):
            report.append("## Executive Summary\n")
            report.append(data['summary'] + "\n")
        
        # Table of Contents
        report.append("## Table of Contents\n")
        sections = ['Introduction', 'Methodology', 'Literature Review', 
                   'Findings', 'Analysis', 'Conclusions', 'References']
        for i, section in enumerate(sections, 1):
            report.append(f"{i}. [{section}](#{section.lower().replace(' ', '-')})")
        report.append("")
        
        # Introduction
        report.append("## Introduction\n")
        report.append(data.get('introduction', 'This research report presents a comprehensive analysis of the collected literature and findings.'))
        report.append("")
        
        # Methodology
        report.append("## Methodology\n")
        report.append("### Data Collection\n")
        report.append(f"- **Documents Analyzed**: {data.get('doc_count', 0)}")
        report.append(f"- **Search Queries**: {', '.join(data.get('queries', []))}")
        report.append(f"- **Date Range**: {data.get('date_range', 'All available')}")
        report.append("")
        
        # Literature Review
        report.append("## Literature Review\n")
        
        if data.get('documents'):
            # Group by year or type
            docs_by_year = defaultdict(list)
            for doc in data['documents']:
                year = doc.get('metadata', {}).get('year', 'Unknown')
                docs_by_year[year].append(doc)
            
            for year in sorted(docs_by_year.keys(), reverse=True):
                if year != 'Unknown':
                    report.append(f"### {year}\n")
                    for doc in docs_by_year[year]:
                        report.append(f"**{doc.get('title', doc.get('filename', 'Untitled'))}**")
                        if doc.get('authors'):
                            report.append(f"*{doc['authors']}*")
                        if doc.get('summary'):
                            report.append(f"\n{doc['summary']}\n")
                        report.append("")
        
        # Findings
        report.append("## Findings\n")
        
        # Key themes
        if data.get('themes'):
            report.append("### Key Themes\n")
            for theme, count in data['themes'].items():
                report.append(f"- **{theme}**: {count} occurrences")
            report.append("")
        
        # Statistics
        if data.get('statistics'):
            report.append("### Statistical Analysis\n")
            stats = data['statistics']
            report.append(f"- Total documents: {stats.get('total_docs', 0)}")
            report.append(f"- Average document length: {stats.get('avg_length', 0):.0f} words")
            report.append(f"- Total citations found: {stats.get('total_citations', 0)}")
            report.append("")
        
        # Analysis
        report.append("## Analysis\n")
        report.append(data.get('analysis', 'Detailed analysis of the findings...'))
        report.append("")
        
        # Conclusions
        report.append("## Conclusions\n")
        report.append(data.get('conclusions', 'Based on the comprehensive analysis...'))
        report.append("")
        
        # References
        report.append("## References\n")
        if data.get('references'):
            for i, ref in enumerate(data['references'], 1):
                report.append(f"{i}. {ref}")
        report.append("")
        
        # Appendices
        if data.get('appendices'):
            report.append("## Appendices\n")
            for title, content in data['appendices'].items():
                report.append(f"### {title}\n")
                report.append(content)
                report.append("")
        
        return '\n'.join(report)
    
    def _generate_html(self, data: Dict) -> str:
        """Generate HTML report"""
        # Convert markdown to HTML
        md_report = self._generate_markdown(data)
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{data.get('title', 'Research Report')}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        .container {{
            background: white;
            padding: 40px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1, h2, h3 {{
            color: #2c3e50;
        }}
        h1 {{
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            margin-top: 30px;
            border-bottom: 1px solid #ecf0f1;
            padding-bottom: 5px;
        }}
        code {{
            background: #f8f9fa;
            padding: 2px 5px;
            border-radius: 3px;
        }}
        blockquote {{
            border-left: 4px solid #3498db;
            padding-left: 20px;
            margin-left: 0;
            color: #666;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background: #3498db;
            color: white;
        }}
        tr:nth-child(even) {{
            background: #f8f9fa;
        }}
        .metadata {{
            background: #ecf0f1;
            padding: 15px;
            border-radius: 5px;
            margin: 10px 0;
        }}
        .summary {{
            background: #e8f4f8;
            padding: 20px;
            border-radius: 5px;
            margin: 20px 0;
            border-left: 5px solid #3498db;
        }}
        @media print {{
            body {{
                background: white;
            }}
            .container {{
                box-shadow: none;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
"""
        
        # Convert Markdown to HTML (simplified)
        if MARKDOWN_AVAILABLE:
            import markdown
            html_content = markdown.markdown(
                md_report, 
                extensions=['tables', 'fenced_code', 'nl2br', 'toc']
            )
        else:
            # Basic conversion
            html_content = md_report.replace('\n', '<br>\n')
            html_content = re.sub(r'# (.*)', r'<h1>\1</h1>', html_content)
            html_content = re.sub(r'## (.*)', r'<h2>\1</h2>', html_content)
            html_content = re.sub(r'### (.*)', r'<h3>\1</h3>', html_content)
            html_content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html_content)
            html_content = re.sub(r'\*(.*?)\*', r'<em>\1</em>', html_content)
        
        html += html_content
        html += """
    </div>
</body>
</html>"""
        
        return html
    
    def _generate_pdf(self, data: Dict) -> bytes:
        """Generate PDF report (requires additional libraries)"""
        # This would use reportlab or weasyprint
        # For now, generate HTML and note that PDF conversion is needed
        html = self._generate_html(data)
        
        # Note: Actual PDF generation would require:
        # from weasyprint import HTML
        # pdf = HTML(string=html).write_pdf()
        
        return html.encode('utf-8')
    
    def _generate_docx(self, data: Dict) -> bytes:
        """Generate DOCX report"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not installed")
        
        doc = DocxDocument()
        
        # Title
        doc.add_heading(data.get('title', 'Research Report'), 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Executive Summary
        if data.get('summary'):
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(data['summary'])
        
        # Add sections
        sections = {
            'Introduction': data.get('introduction', ''),
            'Methodology': data.get('methodology', ''),
            'Findings': data.get('findings', ''),
            'Conclusions': data.get('conclusions', '')
        }
        
        for section_title, content in sections.items():
            if content:
                doc.add_heading(section_title, level=1)
                doc.add_paragraph(content)
        
        # Save to bytes
        from io import BytesIO
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.read()
    
    def _generate_latex(self, data: Dict) -> str:
        """Generate LaTeX report"""
        latex = r"""\documentclass[12pt,a4paper]{article}
\usepackage[utf8]{inputenc}
\usepackage{graphicx}
\usepackage{hyperref}
\usepackage{biblatex}

\title{""" + data.get('title', 'Research Report') + r"""}
\author{Advanced Research System}
\date{\today}

\begin{document}

\maketitle
\tableofcontents
\newpage

"""
        
        # Add sections
        if data.get('summary'):
            latex += r"\section{Executive Summary}" + "\n"
            latex += data['summary'] + "\n\n"
        
        # Introduction
        latex += r"\section{Introduction}" + "\n"
        latex += data.get('introduction', 'Introduction text...') + "\n\n"
        
        # Continue with other sections...
        
        latex += r"\end{document}"
        
        return latex

class AdvancedResearchGUI:
    """Advanced GUI for the research system"""
    
    def __init__(self):
        if not GUI_AVAILABLE:
            raise Exception("Tkinter not available")
        
        self.root = tk.Tk()
        self.root.title("Advanced Research System")
        self.root.geometry("1400x900")
        
        # Initialize components
        self.config = DEFAULT_CONFIG.copy()
        self.db = DatabaseManager(self.config['database_path'])
        self.processor = DocumentProcessor(self.config, self.db)
        self.search_engine = AdvancedSearchEngine(self.config, self.db)
        self.web_engine = WebResearchEngine(self.config, self.db)
        self.summarizer = SummarizationEngine(self.config)
        self.report_gen = ReportGenerator(self.config)
        
        # Current state
        self.current_documents = []
        self.search_results = []
        
        # Setup UI
        self.setup_ui()
        self.apply_theme()
    
    def setup_ui(self):
        """Setup the main UI components"""
        # Create menu bar
        self.create_menu()
        
        # Create main notebook
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Create tabs
        self.create_library_tab()
        self.create_search_tab()
        self.create_analysis_tab()
        self.create_web_research_tab()
        self.create_report_tab()
        self.create_settings_tab()
        
        # Status bar
        self.status_bar = ttk.Label(self.root, text="Ready", relief=tk.SUNKEN)
        self.status_bar.pack(side='bottom', fill='x')
    
    def create_menu(self):
        """Create application menu"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # File menu
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Import Documents...", command=self.import_documents)
        file_menu.add_command(label="Import Folder...", command=self.import_folder)
        file_menu.add_separator()
        file_menu.add_command(label="Export Results...", command=self.export_results)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)
        
        # Edit menu
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Edit", menu=edit_menu)
        edit_menu.add_command(label="Preferences...", command=self.show_preferences)
        
        # Tools menu
        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Tools", menu=tools_menu)
        tools_menu.add_command(label="Batch Process...", command=self.batch_process)
        tools_menu.add_command(label="Citation Manager...", command=self.citation_manager)
        tools_menu.add_command(label="Web Server...", command=self.start_web_server)
        
        # Help menu
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="Documentation", command=self.show_help)
        help_menu.add_command(label="About", command=self.show_about)
    
    def create_library_tab(self):
        """Create document library tab"""
        library_frame = ttk.Frame(self.notebook)
        self.notebook.add(library_frame, text="Document Library")
        
        # Toolbar
        toolbar = ttk.Frame(library_frame)
        toolbar.pack(fill='x', padx=5, pady=5)
        
        ttk.Button(toolbar, text="Add Documents", 
                  command=self.import_documents).pack(side='left', padx=2)
        ttk.Button(toolbar, text="Add Folder", 
                  command=self.import_folder).pack(side='left', padx=2)
        ttk.Button(toolbar, text="Remove Selected", 
                  command=self.remove_documents).pack(side='left', padx=2)
        ttk.Button(toolbar, text="Refresh", 
                  command=self.refresh_library).pack(side='left', padx=2)
        
        # Search bar
        search_frame = ttk.Frame(library_frame)
        search_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(search_frame, text="Filter:").pack(side='left')
        self.library_filter = ttk.Entry(search_frame, width=30)
        self.library_filter.pack(side='left', padx=5, fill='x', expand=True)
        self.library_filter.bind('<KeyRelease>', self.filter_library)
        
        # Document tree
        tree_frame = ttk.Frame(library_frame)
        tree_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Create treeview with columns
        columns = ('Type', 'Size', 'Modified', 'Status')
        self.doc_tree = ttk.Treeview(tree_frame, columns=columns, show='tree headings')
        
        # Configure columns
        self.doc_tree.column('#0', width=400, minwidth=200)
        self.doc_tree.column('Type', width=100)
        self.doc_tree.column('Size', width=100)
        self.doc_tree.column('Modified', width=150)
        self.doc_tree.column('Status', width=100)
        
        # Headings
        self.doc_tree.heading('#0', text='Document')
        self.doc_tree.heading('Type', text='Type')
        self.doc_tree.heading('Size', text='Size')
        self.doc_tree.heading('Modified', text='Modified')
        self.doc_tree.heading('Status', text='Status')
        
        # Scrollbars
        vsb = ttk.Scrollbar(tree_frame, orient='vertical', command=self.doc_tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient='horizontal', command=self.doc_tree.xview)
        self.doc_tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Pack tree and scrollbars
        self.doc_tree.grid(row=0, column=0, sticky='nsew')
        vsb.grid(row=0, column=1, sticky='ns')
        hsb.grid(row=1, column=0, sticky='ew')
        
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)
        
        # Document preview
        preview_frame = ttk.LabelFrame(library_frame, text="Preview", height=200)
        preview_frame.pack(fill='x', padx=5, pady=5)
        preview_frame.pack_propagate(False)
        
        self.preview_text = scrolledtext.ScrolledText(preview_frame, height=8, wrap='word')
        self.preview_text.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Bind events
        self.doc_tree.bind('<<TreeviewSelect>>', self.on_document_select)
        self.doc_tree.bind('<Double-1>', self.open_document)
        
        # Load initial library
        self.refresh_library()
    
    def create_search_tab(self):
        """Create search tab"""
        search_frame = ttk.Frame(self.notebook)
        self.notebook.add(search_frame, text="Search")
        
        # Search input area
        input_frame = ttk.LabelFrame(search_frame, text="Search Query", padding=10)
        input_frame.pack(fill='x', padx=10, pady=5)
        
        # Search type
        type_frame = ttk.Frame(input_frame)
        type_frame.pack(fill='x', pady=5)
        
        ttk.Label(type_frame, text="Search Type:").pack(side='left')
        self.search_type = ttk.Combobox(type_frame, values=['Basic', 'Advanced', 'Semantic'], 
                                       state='readonly', width=15)
        self.search_type.set('Advanced')
        self.search_type.pack(side='left', padx=5)
        
        # Search query
        self.search_query = scrolledtext.ScrolledText(input_frame, height=3, wrap='word')
        self.search_query.pack(fill='x', pady=5)
        
        # Search options
        options_frame = ttk.Frame(input_frame)
        options_frame.pack(fill='x')
        
        self.case_sensitive = tk.BooleanVar()
        ttk.Checkbutton(options_frame, text="Case Sensitive", 
                       variable=self.case_sensitive).pack(side='left', padx=5)
        
        self.whole_words = tk.BooleanVar()
        ttk.Checkbutton(options_frame, text="Whole Words", 
                       variable=self.whole_words).pack(side='left', padx=5)
        
        self.search_metadata = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="Search Metadata", 
                       variable=self.search_metadata).pack(side='left', padx=5)
        
        # Search button
        ttk.Button(input_frame, text="Search", command=self.perform_search,
                  style='Accent.TButton').pack(pady=5)
        
        # Results area
        results_frame = ttk.LabelFrame(search_frame, text="Search Results", padding=10)
        results_frame.pack(fill='both', expand=True, padx=10, pady=5)
        
        # Results tree
        columns = ('Document', 'Page', 'Score', 'Preview')
        self.results_tree = ttk.Treeview(results_frame, columns=columns, show='headings')
        
        # Configure columns
        self.results_tree.column('Document', width=300)
        self.results_tree.column('Page', width=60)
        self.results_tree.column('Score', width=80)
        self.results_tree.column('Preview', width=400)
        
        # Headings
        for col in columns:
            self.results_tree.heading(col, text=col)
        
        # Scrollbar
        results_scroll = ttk.Scrollbar(results_frame, orient='vertical', 
                                     command=self.results_tree.yview)
        self.results_tree.configure(yscrollcommand=results_scroll.set)
        
        # Pack
        self.results_tree.pack(side='left', fill='both', expand=True)
        results_scroll.pack(side='right', fill='y')
        
        # Bind events
        self.results_tree.bind('<Double-1>', self.open_search_result)
    
    def create_analysis_tab(self):
        """Create analysis tab"""
        analysis_frame = ttk.Frame(self.notebook)
        self.notebook.add(analysis_frame, text="Analysis")
        
        # Analysis options
        options_frame = ttk.LabelFrame(analysis_frame, text="Analysis Options", padding=10)
        options_frame.pack(fill='x', padx=10, pady=5)
        
        # Analysis type selection
        self.analysis_types = {
            'statistics': tk.BooleanVar(value=True),
            'themes': tk.BooleanVar(value=True),
            'citations': tk.BooleanVar(value=True),
            'timeline': tk.BooleanVar(value=False),
            'network': tk.BooleanVar(value=False)
        }
        
        for name, var in self.analysis_types.items():
            ttk.Checkbutton(options_frame, text=name.title(), 
                           variable=var).pack(side='left', padx=5)
        
        # Run analysis button
        ttk.Button(options_frame, text="Run Analysis", 
                  command=self.run_analysis, style='Accent.TButton').pack(pady=5)
        
        # Results notebook
        analysis_notebook = ttk.Notebook(analysis_frame)
        analysis_notebook.pack(fill='both', expand=True, padx=10, pady=5)
        
        # Statistics tab
        stats_frame = ttk.Frame(analysis_notebook)
        analysis_notebook.add(stats_frame, text="Statistics")
        
        self.stats_text = scrolledtext.ScrolledText(stats_frame, wrap='word')
        self.stats_text.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Themes tab
        themes_frame = ttk.Frame(analysis_notebook)
        analysis_notebook.add(themes_frame, text="Themes")
        
        self.themes_tree = ttk.Treeview(themes_frame, columns=('Count', 'Documents'), 
                                       show='tree headings')
        self.themes_tree.column('#0', width=300)
        self.themes_tree.column('Count', width=100)
        self.themes_tree.column('Documents', width=400)
        self.themes_tree.heading('#0', text='Theme')
        self.themes_tree.heading('Count', text='Count')
        self.themes_tree.heading('Documents', text='Documents')
        self.themes_tree.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Citations tab
        citations_frame = ttk.Frame(analysis_notebook)
        analysis_notebook.add(citations_frame, text="Citations")
        
        self.citations_text = scrolledtext.ScrolledText(citations_frame, wrap='word')
        self.citations_text.pack(fill='both', expand=True, padx=5, pady=5)
    
    def create_web_research_tab(self):
        """Create web research tab"""
        web_frame = ttk.Frame(self.notebook)
        self.notebook.add(web_frame, text="Web Research")
        
        # Search area
        search_frame = ttk.LabelFrame(web_frame, text="Web Search", padding=10)
        search_frame.pack(fill='x', padx=10, pady=5)
        
        # Query input
        ttk.Label(search_frame, text="Search Query:").pack(anchor='w')
        self.web_query = ttk.Entry(search_frame, width=50)
        self.web_query.pack(fill='x', pady=5)
        
        # Options
        options_frame = ttk.Frame(search_frame)
        options_frame.pack(fill='x')
        
        ttk.Label(options_frame, text="Max Results:").pack(side='left')
        self.max_results = ttk.Spinbox(options_frame, from_=10, to=100, 
                                      increment=10, width=10)
        self.max_results.set(30)
        self.max_results.pack(side='left', padx=5)
        
        self.enable_crawl = tk.BooleanVar()
        ttk.Checkbutton(options_frame, text="Enable Deep Crawl", 
                       variable=self.enable_crawl).pack(side='left', padx=10)
        
        # Search button
        ttk.Button(search_frame, text="Search Web", 
                  command=self.search_web, style='Accent.TButton').pack(pady=5)
        
        # Results
        results_frame = ttk.LabelFrame(web_frame, text="Web Results", padding=10)
        results_frame.pack(fill='both', expand=True, padx=10, pady=5)
        
        # Results listbox
        self.web_results = tk.Listbox(results_frame, height=10)
        web_scroll = ttk.Scrollbar(results_frame, orient='vertical', 
                                  command=self.web_results.yview)
        self.web_results.configure(yscrollcommand=web_scroll.set)
        
        self.web_results.pack(side='left', fill='both', expand=True)
        web_scroll.pack(side='right', fill='y')
        
        # Preview
        preview_frame = ttk.LabelFrame(web_frame, text="Preview", padding=10)
        preview_frame.pack(fill='both', expand=True, padx=10, pady=5)
        
        self.web_preview = scrolledtext.ScrolledText(preview_frame, wrap='word', height=10)
        self.web_preview.pack(fill='both', expand=True)
        
        # Bind events
        self.web_results.bind('<<ListboxSelect>>', self.preview_web_result)
    
    def create_report_tab(self):
        """Create report generation tab"""
        report_frame = ttk.Frame(self.notebook)
        self.notebook.add(report_frame, text="Report")
        
        # Report options
        options_frame = ttk.LabelFrame(report_frame, text="Report Options", padding=10)
        options_frame.pack(fill='x', padx=10, pady=5)
        
        # Title
        ttk.Label(options_frame, text="Report Title:").grid(row=0, column=0, sticky='w', pady=5)
        self.report_title = ttk.Entry(options_frame, width=50)
        self.report_title.grid(row=0, column=1, pady=5, padx=5)
        self.report_title.insert(0, "Research Report")
        
        # Format
        ttk.Label(options_frame, text="Format:").grid(row=1, column=0, sticky='w', pady=5)
        self.report_format = ttk.Combobox(options_frame, 
                                         values=['Markdown', 'HTML', 'PDF', 'DOCX', 'LaTeX'],
                                         state='readonly', width=20)
        self.report_format.set('Markdown')
        self.report_format.grid(row=1, column=1, sticky='w', pady=5, padx=5)
        
        # Sections to include
        sections_frame = ttk.LabelFrame(options_frame, text="Include Sections", padding=5)
        sections_frame.grid(row=2, column=0, columnspan=2, pady=10, sticky='ew')
        
        self.report_sections = {
            'summary': tk.BooleanVar(value=True),
            'introduction': tk.BooleanVar(value=True),
            'methodology': tk.BooleanVar(value=True),
            'literature_review': tk.BooleanVar(value=True),
            'findings': tk.BooleanVar(value=True),
            'analysis': tk.BooleanVar(value=True),
            'conclusions': tk.BooleanVar(value=True),
            'references': tk.BooleanVar(value=True)
        }
        
        for i, (name, var) in enumerate(self.report_sections.items()):
            ttk.Checkbutton(sections_frame, text=name.replace('_', ' ').title(), 
                           variable=var).grid(row=i//4, column=i%4, padx=5, pady=2, sticky='w')
        
        # Generate button
        ttk.Button(options_frame, text="Generate Report", 
                  command=self.generate_report, style='Accent.TButton').grid(
                      row=3, column=0, columnspan=2, pady=10)
        
        # Preview area
        preview_frame = ttk.LabelFrame(report_frame, text="Report Preview", padding=10)
        preview_frame.pack(fill='both', expand=True, padx=10, pady=5)
        
        self.report_preview = scrolledtext.ScrolledText(preview_frame, wrap='word')
        self.report_preview.pack(fill='both', expand=True)
        
        # Export button
        ttk.Button(preview_frame, text="Export Report", 
                  command=self.export_report).pack(pady=5)
    
    def create_settings_tab(self):
        """Create settings tab"""
        settings_frame = ttk.Frame(self.notebook)
        self.notebook.add(settings_frame, text="Settings")
        
        # Create scrollable frame
        canvas = tk.Canvas(settings_frame)
        scrollbar = ttk.Scrollbar(settings_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # General settings
        general_frame = ttk.LabelFrame(scrollable_frame, text="General Settings", padding=10)
        general_frame.pack(fill='x', padx=10, pady=5)
        
        # Output directory
        ttk.Label(general_frame, text="Output Directory:").grid(row=0, column=0, sticky='w', pady=5)
        self.output_dir_var = tk.StringVar(value=self.config['output_dir'])
        ttk.Entry(general_frame, textvariable=self.output_dir_var, width=40).grid(
            row=0, column=1, pady=5, padx=5)
        ttk.Button(general_frame, text="Browse...", 
                  command=self.browse_output_dir).grid(row=0, column=2, pady=5)
        
        # Processing settings
        processing_frame = ttk.LabelFrame(scrollable_frame, text="Processing Settings", padding=10)
        processing_frame.pack(fill='x', padx=10, pady=5)
        
        # Max threads
        ttk.Label(processing_frame, text="Max Threads:").grid(row=0, column=0, sticky='w', pady=5)
        self.max_threads_var = tk.IntVar(value=self.config['max_threads'])
        ttk.Spinbox(processing_frame, from_=1, to=16, textvariable=self.max_threads_var, 
                   width=10).grid(row=0, column=1, sticky='w', pady=5, padx=5)
        
        # OCR
        self.ocr_enabled_var = tk.BooleanVar(value=self.config['ocr_enabled'])
        ttk.Checkbutton(processing_frame, text="Enable OCR for scanned PDFs", 
                       variable=self.ocr_enabled_var).grid(row=1, column=0, columnspan=2, 
                                                           sticky='w', pady=5)
        
        # AI settings
        ai_frame = ttk.LabelFrame(scrollable_frame, text="AI/Summarization Settings", padding=10)
        ai_frame.pack(fill='x', padx=10, pady=5)
        
        # Summarization method
        ttk.Label(ai_frame, text="Summarization Method:").grid(row=0, column=0, sticky='w', pady=5)
        self.summary_method_var = tk.StringVar(value=self.config['summary_method'])
        ttk.Combobox(ai_frame, textvariable=self.summary_method_var,
                    values=['textrank', 'frequency', 'ollama', 'openai', 'hybrid'],
                    state='readonly', width=20).grid(row=0, column=1, sticky='w', pady=5, padx=5)
        
        # Ollama URL
        ttk.Label(ai_frame, text="Ollama URL:").grid(row=1, column=0, sticky='w', pady=5)
        self.ollama_url_var = tk.StringVar(value=self.config['ollama_url'])
        ttk.Entry(ai_frame, textvariable=self.ollama_url_var, width=40).grid(
            row=1, column=1, pady=5, padx=5)
        
        # OpenAI API Key
        ttk.Label(ai_frame, text="OpenAI API Key:").grid(row=2, column=0, sticky='w', pady=5)
        self.openai_key_var = tk.StringVar(value=self.config['openai_api_key'])
        ttk.Entry(ai_frame, textvariable=self.openai_key_var, width=40, show='*').grid(
            row=2, column=1, pady=5, padx=5)
        
        # Save settings button
        ttk.Button(scrollable_frame, text="Save Settings", 
                  command=self.save_settings, style='Accent.TButton').pack(pady=10)
        
        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
    
    def apply_theme(self):
        """Apply modern theme to the application"""
        style = ttk.Style()
        
        # Configure styles
        style.configure('TNotebook', background='#f0f0f0')
        style.configure('TNotebook.Tab', padding=[20, 10], font=('Arial', 10))
        style.configure('TFrame', background='#ffffff')
        style.configure('TLabelframe', background='#ffffff')
        style.configure('TLabelframe.Label', font=('Arial', 10, 'bold'))
        
        # Accent button
        style.configure('Accent.TButton', font=('Arial', 10, 'bold'))
        style.map('Accent.TButton',
                 background=[('active', '#0056b3'), ('!active', '#007bff')],
                 foreground=[('active', 'white'), ('!active', 'white')])
    
    # Event handlers and methods
    def import_documents(self):
        """Import documents into the library"""
        files = filedialog.askopenfilenames(
            title="Select Documents",
            filetypes=[
                ("All Supported", "*.pdf;*.docx;*.txt;*.md;*.html;*.htm"),
                ("PDF Files", "*.pdf"),
                ("Word Documents", "*.docx"),
                ("Text Files", "*.txt"),
                ("Markdown Files", "*.md"),
                ("HTML Files", "*.html;*.htm"),
                ("All Files", "*.*")
            ]
        )
        
        if files:
            self.process_files(files)
    
    def import_folder(self):
        """Import all documents from a folder"""
        folder = filedialog.askdirectory(title="Select Folder")
        if not folder:
            return
        
        # Find all supported files
        files = []
        for ext in self.config['supported_formats']:
            files.extend(Path(folder).rglob(f"*{ext}"))
        
        if files:
            self.process_files([str(f) for f in files])
        else:
            messagebox.showinfo("No Files", "No supported files found in the selected folder")
    
    def process_files(self, files):
        """Process imported files"""
        progress = ttk.Progressbar(self.root, mode='determinate', maximum=len(files))
        progress.place(relx=0.5, rely=0.5, anchor='center')
        
        processed = 0
        errors = []
        
        for file in files:
            try:
                self.status_bar.config(text=f"Processing: {os.path.basename(file)}")
                self.root.update()
                
                result = self.processor.process_document(file)
                
                if result.get('error'):
                    errors.append(f"{os.path.basename(file)}: {result['error']}")
                else:
                    processed += 1
                
                progress['value'] = files.index(file) + 1
                self.root.update()
                
            except Exception as e:
                errors.append(f"{os.path.basename(file)}: {str(e)}")
        
        progress.destroy()
        
        # Show results
        if errors:
            error_msg = "\n".join(errors[:10])  # Show first 10 errors
            if len(errors) > 10:
                error_msg += f"\n... and {len(errors) - 10} more errors"
            messagebox.showwarning("Processing Errors", error_msg)
        
        messagebox.showinfo("Import Complete", 
                          f"Successfully processed {processed} out of {len(files)} files")
        
        self.refresh_library()
        self.status_bar.config(text="Ready")
    
    def refresh_library(self):
        """Refresh the document library display"""
        # Clear tree
        for item in self.doc_tree.get_children():
            self.doc_tree.delete(item)
        
        # Get documents from database
        with sqlite3.connect(self.db.db_path) as conn:
            cursor = conn.cursor()
            
            # First check what columns exist
            cursor.execute("PRAGMA table_info(documents)")
            columns = [col[1] for col in cursor.fetchall()]
            
            # Build query based on available columns
            select_columns = ['id', 'filename', 'filepath', 'filetype']
            
            # Add optional columns if they exist
            if 'filesize' in columns:
                select_columns.append('filesize')
            else:
                select_columns.append('0 as filesize')
                
            if 'date_modified' in columns:
                select_columns.append('date_modified')
            else:
                select_columns.append('NULL as date_modified')
                
            if 'processed' in columns:
                select_columns.append('processed')
            else:
                select_columns.append('1 as processed')
            
            query = f"""
                SELECT {', '.join(select_columns)}
                FROM documents
                ORDER BY filename
            """
            
            cursor.execute(query)
            
            for row in cursor.fetchall():
                doc_id, filename, filepath, filetype = row[:4]
                filesize = row[4] if len(row) > 4 else 0
                date_modified = row[5] if len(row) > 5 else None
                processed = row[6] if len(row) > 6 else True
                
                # Format values
                size_str = self.format_size(filesize) if filesize else "Unknown"
                date_str = date_modified[:10] if date_modified else "Unknown"
                status = "Ready" if processed else "Processing"
                
                # Insert into tree
                self.doc_tree.insert('', 'end', values=(filetype, size_str, date_str, status),
                                   text=filename, tags=(doc_id,))
    
    def format_size(self, bytes):
        """Format file size"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if bytes < 1024.0:
                return f"{bytes:.1f} {unit}"
            bytes /= 1024.0
        return f"{bytes:.1f} TB"
    
    def on_document_select(self, event):
        """Handle document selection"""
        selection = self.doc_tree.selection()
        if not selection:
            return
        
        # Get document ID from tags
        item = self.doc_tree.item(selection[0])
        doc_id = item['tags'][0] if item['tags'] else None
        
        if doc_id:
            # Load preview
            with sqlite3.connect(self.db.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT full_text FROM documents WHERE id = ?", (doc_id,))
                result = cursor.fetchone()
                
                if result and result[0]:
                    # Show first 1000 characters
                    preview = result[0][:1000]
                    if len(result[0]) > 1000:
                        preview += "\n\n... (truncated)"
                    
                    self.preview_text.delete(1.0, tk.END)
                    self.preview_text.insert(1.0, preview)
    
    def open_document(self, event):
        """Open selected document"""
        selection = self.doc_tree.selection()
        if not selection:
            return
        
        # Get filepath
        item = self.doc_tree.item(selection[0])
        doc_id = item['tags'][0] if item['tags'] else None
        
        if doc_id:
            with sqlite3.connect(self.db.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT filepath FROM documents WHERE id = ?", (doc_id,))
                result = cursor.fetchone()
                
                if result:
                    filepath = result[0]
                    if os.path.exists(filepath):
                        try:
                            if sys.platform.startswith('darwin'):
                                subprocess.call(['open', filepath])
                            elif sys.platform.startswith('win'):
                                os.startfile(filepath)
                            else:  # linux
                                subprocess.call(['xdg-open', filepath])
                        except Exception as e:
                            messagebox.showerror("Error", f"Could not open file: {str(e)}")
    
    def remove_documents(self):
        """Remove selected documents from library"""
        selection = self.doc_tree.selection()
        if not selection:
            messagebox.showinfo("No Selection", "Please select documents to remove")
            return
        
        if messagebox.askyesno("Confirm", f"Remove {len(selection)} document(s) from library?"):
            doc_ids = []
            for item in selection:
                tags = self.doc_tree.item(item)['tags']
                if tags:
                    doc_ids.append(tags[0])
            
            # Remove from database
            with sqlite3.connect(self.db.db_path) as conn:
                cursor = conn.cursor()
                for doc_id in doc_ids:
                    cursor.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
                    cursor.execute("DELETE FROM metadata WHERE document_id = ?", (doc_id,))
                    cursor.execute("DELETE FROM paragraphs WHERE document_id = ?", (doc_id,))
                conn.commit()
            
            self.refresh_library()
    
    def filter_library(self, event):
        """Filter document library"""
        filter_text = self.library_filter.get().lower()
        
        # Clear and repopulate tree
        for item in self.doc_tree.get_children():
            self.doc_tree.delete(item)
        
        # Get filtered documents
        with sqlite3.connect(self.db.db_path) as conn:
            cursor = conn.cursor()
            
            # Check available columns
            cursor.execute("PRAGMA table_info(documents)")
            columns = [col[1] for col in cursor.fetchall()]
            
            # Build query with available columns
            select_columns = ['id', 'filename', 'filepath', 'filetype']
            
            if 'filesize' in columns:
                select_columns.append('filesize')
            else:
                select_columns.append('0 as filesize')
                
            if 'date_modified' in columns:
                select_columns.append('date_modified')
            else:
                select_columns.append('NULL as date_modified')
                
            if 'processed' in columns:
                select_columns.append('processed')
            else:
                select_columns.append('1 as processed')
            
            if filter_text:
                query = f"""
                    SELECT {', '.join(select_columns)}
                    FROM documents
                    WHERE LOWER(filename) LIKE ?
                    ORDER BY filename
                """
                cursor.execute(query, (f"%{filter_text}%",))
            else:
                query = f"""
                    SELECT {', '.join(select_columns)}
                    FROM documents
                    ORDER BY filename
                """
                cursor.execute(query)
            
            for row in cursor.fetchall():
                doc_id, filename, filepath, filetype = row[:4]
                filesize = row[4] if len(row) > 4 else 0
                date_modified = row[5] if len(row) > 5 else None
                processed = row[6] if len(row) > 6 else True
                
                size_str = self.format_size(filesize) if filesize else "Unknown"
                date_str = date_modified[:10] if date_modified else "Unknown"
                status = "Ready" if processed else "Processing"
                
                self.doc_tree.insert('', 'end', values=(filetype, size_str, date_str, status),
                                   text=filename, tags=(doc_id,))
    
    def perform_search(self):
        """Perform document search"""
        query = self.search_query.get(1.0, tk.END).strip()
        if not query:
            messagebox.showinfo("No Query", "Please enter a search query")
            return
        
        # Update config
        self.config['case_sensitive'] = self.case_sensitive.get()
        self.config['whole_words'] = self.whole_words.get()
        
        # Get search method
        method = self.search_type.get().lower()
        
        # Perform search
        self.status_bar.config(text="Searching...")
        self.root.update()
        
        try:
            results = self.search_engine.search(query, method)
            
            # Clear previous results
            for item in self.results_tree.get_children():
                self.results_tree.delete(item)
            
            # Display results
            for result in results:
                preview = ""
                if result.get('paragraph_content'):
                    preview = result['paragraph_content'][:200] + "..."
                
                self.results_tree.insert('', 'end', values=(
                    result['filename'],
                    result.get('page_num', 'N/A'),
                    f"{result['relevance_score']:.2f}",
                    preview
                ))
            
            self.status_bar.config(text=f"Found {len(results)} results")
            
        except Exception as e:
            messagebox.showerror("Search Error", str(e))
            self.status_bar.config(text="Search failed")
    
    def open_search_result(self, event):
        """Open selected search result"""
        selection = self.results_tree.selection()
        if selection:
            item = self.results_tree.item(selection[0])
            filename = item['values'][0]
            
            # Find and open the document
            with sqlite3.connect(self.db.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT filepath FROM documents WHERE filename = ?", (filename,))
                result = cursor.fetchone()
                
                if result:
                    filepath = result[0]
                    if os.path.exists(filepath):
                        try:
                            if sys.platform.startswith('darwin'):
                                subprocess.call(['open', filepath])
                            elif sys.platform.startswith('win'):
                                os.startfile(filepath)
                            else:
                                subprocess.call(['xdg-open', filepath])
                        except Exception as e:
                            messagebox.showerror("Error", f"Could not open file: {str(e)}")
    
    def run_analysis(self):
        """Run selected analyses"""
        self.status_bar.config(text="Running analysis...")
        self.root.update()
        
        try:
            # Get all documents
            with sqlite3.connect(self.db.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT COUNT(*) FROM documents WHERE processed = 1")
                doc_count = cursor.fetchone()[0]
            
            if doc_count == 0:
                messagebox.showinfo("No Documents", "No processed documents available for analysis")
                return
            
            # Run statistics
            if self.analysis_types['statistics'].get():
                stats = self.calculate_statistics()
                self.display_statistics(stats)
            
            # Extract themes
            if self.analysis_types['themes'].get():
                themes = self.extract_themes()
                self.display_themes(themes)
            
            # Analyze citations
            if self.analysis_types['citations'].get():
                citations = self.analyze_citations()
                self.display_citations(citations)
            
            self.status_bar.config(text="Analysis complete")
            
        except Exception as e:
            messagebox.showerror("Analysis Error", str(e))
            self.status_bar.config(text="Analysis failed")
    
    def calculate_statistics(self):
        """Calculate document statistics"""
        stats = {}
        
        with sqlite3.connect(self.db.db_path) as conn:
            cursor = conn.cursor()
            
            # Document count
            cursor.execute("SELECT COUNT(*) FROM documents WHERE processed = 1")
            stats['total_documents'] = cursor.fetchone()[0]
            
            # Total size
            cursor.execute("SELECT SUM(filesize) FROM documents")
            total_size = cursor.fetchone()[0] or 0
            stats['total_size'] = self.format_size(total_size)
            
            # Document types
            cursor.execute("""
                SELECT filetype, COUNT(*) 
                FROM documents 
                GROUP BY filetype
            """)
            stats['document_types'] = dict(cursor.fetchall())
            
            # Average document length
            cursor.execute("""
                SELECT AVG(LENGTH(full_text)) 
                FROM documents 
                WHERE full_text IS NOT NULL
            """)
            avg_length = cursor.fetchone()[0] or 0
            stats['avg_document_length'] = int(avg_length)
            
            # Total paragraphs
            cursor.execute("SELECT COUNT(*) FROM paragraphs")
            stats['total_paragraphs'] = cursor.fetchone()[0]
            
            # Metadata statistics
            cursor.execute("""
                SELECT key, COUNT(DISTINCT document_id) 
                FROM metadata 
                GROUP BY key 
                ORDER BY COUNT(DISTINCT document_id) DESC
            """)
            stats['metadata_fields'] = dict(cursor.fetchall())
        
        return stats
    
    def display_statistics(self, stats):
        """Display statistics in the UI"""
        self.stats_text.delete(1.0, tk.END)
        
        text = "DOCUMENT STATISTICS\n" + "="*50 + "\n\n"
        
        text += f"Total Documents: {stats['total_documents']}\n"
        text += f"Total Size: {stats['total_size']}\n"
        text += f"Average Document Length: {stats['avg_document_length']:,} characters\n"
        text += f"Total Paragraphs: {stats['total_paragraphs']:,}\n\n"
        
        text += "Document Types:\n"
        for filetype, count in stats['document_types'].items():
            text += f"  {filetype}: {count}\n"
        
        text += "\nMetadata Fields:\n"
        for field, count in list(stats['metadata_fields'].items())[:10]:
            text += f"  {field}: {count} documents\n"
        
        self.stats_text.insert(1.0, text)
    
    def extract_themes(self):
        """Extract themes from documents"""
        themes = Counter()
        
        # Simple keyword extraction (would be enhanced with NLP)
        with sqlite3.connect(self.db.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT full_text 
                FROM documents 
                WHERE full_text IS NOT NULL
            """)
            
            for row in cursor.fetchall():
                text = row[0].lower()
                
                # Extract keywords (simplified)
                words = re.findall(r'\b\w{4,}\b', text)
                
                # Filter common words
                stopwords = {'this', 'that', 'with', 'from', 'have', 'been', 
                           'were', 'their', 'would', 'could', 'should'}
                
                keywords = [w for w in words if w not in stopwords]
                
                # Count most common
                word_counts = Counter(keywords)
                themes.update(word_counts.most_common(20))
        
        return themes.most_common(50)
    
    def display_themes(self, themes):
        """Display themes in the UI"""
        # Clear tree
        for item in self.themes_tree.get_children():
            self.themes_tree.delete(item)
        
        # Add themes
        for theme, count in themes:
            self.themes_tree.insert('', 'end', text=theme, 
                                   values=(count, "Multiple documents"))
    
    def analyze_citations(self):
        """Analyze citations in documents"""
        citations = []
        
        # Simple citation extraction (would be enhanced with citation parsing)
        with sqlite3.connect(self.db.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT filename, full_text 
                FROM documents 
                WHERE full_text IS NOT NULL
            """)
            
            for filename, text in cursor.fetchall():
                # Look for common citation patterns
                # Year in parentheses
                year_matches = re.findall(r'\((\d{4})\)', text)
                
                # Author (Year) pattern
                author_year = re.findall(r'([A-Z][a-zA-Z]+(?:\s+(?:and|&)\s+[A-Z][a-zA-Z]+)*)\s*\((\d{4})\)', text)
                
                for match in author_year[:10]:  # Limit to first 10
                    citations.append({
                        'document': filename,
                        'authors': match[0],
                        'year': match[1],
                        'type': 'in-text'
                    })
        
        return citations
    
    def display_citations(self, citations):
        """Display citations in the UI"""
        self.citations_text.delete(1.0, tk.END)
        
        text = "CITATION ANALYSIS\n" + "="*50 + "\n\n"
        
        if not citations:
            text += "No citations found.\n"
        else:
            # Group by document
            by_doc = defaultdict(list)
            for cit in citations:
                by_doc[cit['document']].append(cit)
            
            for doc, cits in by_doc.items():
                text += f"\n{doc}:\n"
                for cit in cits:
                    text += f"  - {cit['authors']} ({cit['year']})\n"
        
        self.citations_text.insert(1.0, text)
    
    def search_web(self):
        """Perform web search"""
        query = self.web_query.get().strip()
        if not query:
            messagebox.showinfo("No Query", "Please enter a search query")
            return
        
        self.status_bar.config(text="Searching web...")
        self.root.update()
        
        try:
            # Clear previous results
            self.web_results.delete(0, tk.END)
            self.web_preview.delete(1.0, tk.END)
            
            # Search
            max_results = int(self.max_results.get())
            results = self.web_engine.search_web(query, max_results)
            
            # Display results
            self.web_search_results = results  # Store for preview
            for i, result in enumerate(results):
                display_text = f"{i+1}. {result['title']}"
                self.web_results.insert(tk.END, display_text)
            
            self.status_bar.config(text=f"Found {len(results)} web results")
            
            # Crawl if enabled
            if self.enable_crawl.get() and results:
                self.status_bar.config(text="Crawling web pages...")
                self.root.update()
                
                # Crawl first few results
                for result in results[:5]:
                    self.web_engine.crawl_url(result['url'], depth=1)
                
                self.status_bar.config(text="Web crawl complete")
            
        except Exception as e:
            messagebox.showerror("Web Search Error", str(e))
            self.status_bar.config(text="Web search failed")
    
    def preview_web_result(self, event):
        """Preview selected web result"""
        selection = self.web_results.curselection()
        if not selection:
            return
        
        index = selection[0]
        if index < len(self.web_search_results):
            result = self.web_search_results[index]
            
            preview = f"Title: {result['title']}\n"
            preview += f"URL: {result['url']}\n\n"
            preview += f"Snippet: {result['snippet']}\n"
            
            self.web_preview.delete(1.0, tk.END)
            self.web_preview.insert(1.0, preview)
    
    def generate_report(self):
        """Generate research report"""
        self.status_bar.config(text="Generating report...")
        self.root.update()
        
        try:
            # Gather data for report
            report_data = {
                'title': self.report_title.get(),
                'timestamp': datetime.now(),
                'queries': []  # Would be populated from search history
            }
            
            # Get document data
            with sqlite3.connect(self.db.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT d.id, d.filename, d.filepath, d.full_text
                    FROM documents d
                    WHERE d.processed = 1
                """)
                
                documents = []
                for doc_id, filename, filepath, full_text in cursor.fetchall():
                    # Get metadata
                    metadata = self.db.get_document_metadata(doc_id)
                    
                    # Generate summary
                    summary = ""
                    if self.config.get('summarization_enabled') and full_text:
                        summary = self.summarizer.summarize(full_text[:5000], length='medium')
                    
                    documents.append({
                        'id': doc_id,
                        'filename': filename,
                        'title': metadata.get('title', filename),
                        'authors': metadata.get('author', metadata.get('authors', '')),
                        'year': metadata.get('year', ''),
                        'metadata': metadata,
                        'summary': summary
                    })
                
                report_data['documents'] = documents
                report_data['doc_count'] = len(documents)
            
            # Add sections based on checkboxes
            if self.report_sections['summary'].get():
                report_data['summary'] = self.generate_executive_summary(documents)
            
            if self.report_sections['introduction'].get():
                report_data['introduction'] = "This comprehensive research report presents an analysis of the collected literature..."
            
            if self.report_sections['methodology'].get():
                report_data['methodology'] = "The research methodology involved systematic collection and analysis of documents..."
            
            # Get statistics
            if self.report_sections['analysis'].get():
                stats = self.calculate_statistics()
                report_data['statistics'] = stats
                
                # Get themes
                themes = self.extract_themes()
                report_data['themes'] = dict(themes[:20])
            
            # Generate report
            format_map = {
                'Markdown': 'markdown',
                'HTML': 'html',
                'PDF': 'pdf',
                'DOCX': 'docx',
                'LaTeX': 'latex'
            }
            
            format = format_map.get(self.report_format.get(), 'markdown')
            report_content = self.report_gen.generate_report(report_data, format)
            
            # Display preview
            self.report_preview.delete(1.0, tk.END)
            if format in ['markdown', 'latex']:
                self.report_preview.insert(1.0, report_content)
            elif format == 'html':
                # Show raw HTML for now
                self.report_preview.insert(1.0, report_content[:2000] + "\n\n... (HTML preview truncated)")
            else:
                self.report_preview.insert(1.0, f"Report generated in {format} format. Click 'Export Report' to save.")
            
            self.current_report = report_content
            self.current_report_format = format
            
            self.status_bar.config(text="Report generated successfully")
            
        except Exception as e:
            messagebox.showerror("Report Generation Error", str(e))
            self.status_bar.config(text="Report generation failed")
    
    def generate_executive_summary(self, documents):
        """Generate executive summary for report"""
        if not documents:
            return "No documents available for summary."
        
        # Combine key information
        summary_parts = []
        
        # Document overview
        summary_parts.append(f"This report analyzes {len(documents)} documents")
        
        # Year range
        years = [int(doc['year']) for doc in documents if doc['year'] and doc['year'].isdigit()]
        if years:
            summary_parts.append(f"spanning from {min(years)} to {max(years)}")
        
        # Key topics (simplified)
        all_titles = ' '.join(doc['title'] for doc in documents if doc['title'])
        common_words = Counter(re.findall(r'\b\w{5,}\b', all_titles.lower())).most_common(5)
        if common_words:
            topics = ', '.join(word[0] for word in common_words)
            summary_parts.append(f"Key topics include: {topics}")
        
        return '. '.join(summary_parts) + '.'
    
    def export_report(self):
        """Export the generated report"""
        if not hasattr(self, 'current_report'):
            messagebox.showinfo("No Report", "Please generate a report first")
            return
        
        # Get file extension
        ext_map = {
            'markdown': '.md',
            'html': '.html',
            'pdf': '.pdf',
            'docx': '.docx',
            'latex': '.tex'
        }
        
        ext = ext_map.get(self.current_report_format, '.txt')
        
        # Ask for save location
        filename = filedialog.asksaveasfilename(
            defaultextension=ext,
            filetypes=[
                (f"{self.current_report_format.upper()} files", f"*{ext}"),
                ("All files", "*.*")
            ]
        )
        
        if filename:
            try:
                if self.current_report_format in ['markdown', 'html', 'latex']:
                    with open(filename, 'w', encoding='utf-8') as f:
                        f.write(self.current_report)
                else:
                    # Binary formats
                    with open(filename, 'wb') as f:
                        f.write(self.current_report)
                
                messagebox.showinfo("Export Complete", f"Report saved to:\n{filename}")
                
            except Exception as e:
                messagebox.showerror("Export Error", str(e))
    
    def export_results(self):
        """Export search/analysis results"""
        export_window = tk.Toplevel(self.root)
        export_window.title("Export Results")
        export_window.geometry("400x300")
        
        # Export options
        ttk.Label(export_window, text="Select export format:", 
                 font=('Arial', 10, 'bold')).pack(pady=10)
        
        format_var = tk.StringVar(value="xlsx")
        formats = [
            ("Excel Spreadsheet (.xlsx)", "xlsx"),
            ("CSV File (.csv)", "csv"),
            ("JSON File (.json)", "json"),
            ("BibTeX File (.bib)", "bibtex")
        ]
        
        for text, value in formats:
            ttk.Radiobutton(export_window, text=text, variable=format_var, 
                           value=value).pack(anchor='w', padx=20, pady=5)
        
        # Export button
        def do_export():
            format = format_var.get()
            
            # Get save location
            ext_map = {'xlsx': '.xlsx', 'csv': '.csv', 'json': '.json', 'bibtex': '.bib'}
            filename = filedialog.asksaveasfilename(
                defaultextension=ext_map[format],
                filetypes=[(f"{format.upper()} files", f"*{ext_map[format]}")]
            )
            
            if filename:
                try:
                    self.export_data(filename, format)
                    messagebox.showinfo("Export Complete", f"Data exported to:\n{filename}")
                    export_window.destroy()
                except Exception as e:
                    messagebox.showerror("Export Error", str(e))
        
        ttk.Button(export_window, text="Export", command=do_export,
                  style='Accent.TButton').pack(pady=20)
    
    def export_data(self, filename, format):
        """Export data in specified format"""
        # Get all document data
        with sqlite3.connect(self.db.db_path) as conn:
            if format == 'xlsx':
                # Use pandas for Excel export
                if not PANDAS_AVAILABLE:
                    raise Exception("pandas not installed")
                
                # Documents sheet
                df_docs = pd.read_sql_query("""
                    SELECT filename, filepath, filetype, filesize, 
                           date_added, date_modified, processed
                    FROM documents
                """, conn)
                
                # Metadata sheet
                df_meta = pd.read_sql_query("""
                    SELECT d.filename, m.key, m.value
                    FROM metadata m
                    JOIN documents d ON m.document_id = d.id
                """, conn)
                
                # Write to Excel with multiple sheets
                with pd.ExcelWriter(filename, engine='openpyxl') as writer:
                    df_docs.to_excel(writer, sheet_name='Documents', index=False)
                    df_meta.to_excel(writer, sheet_name='Metadata', index=False)
                    
                    # Add search results if any
                    if hasattr(self, 'search_results') and self.search_results:
                        results_data = []
                        for r in self.search_results:
                            results_data.append({
                                'filename': r.get('filename'),
                                'page': r.get('page_num'),
                                'score': r.get('relevance_score'),
                                'preview': r.get('paragraph_content', '')[:200]
                            })
                        df_results = pd.DataFrame(results_data)
                        df_results.to_excel(writer, sheet_name='Search Results', index=False)
            
            elif format == 'csv':
                # Export documents to CSV
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT filename, filepath, filetype, filesize, 
                           date_added, date_modified, processed
                    FROM documents
                """)
                
                with open(filename, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow(['Filename', 'Path', 'Type', 'Size', 
                                   'Date Added', 'Date Modified', 'Processed'])
                    writer.writerows(cursor.fetchall())
            
            elif format == 'json':
                # Export all data as JSON
                data = {
                    'export_date': datetime.now().isoformat(),
                    'documents': []
                }
                
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM documents")
                columns = [desc[0] for desc in cursor.description]
                
                for row in cursor.fetchall():
                    doc = dict(zip(columns, row))
                    doc['metadata'] = self.db.get_document_metadata(doc['id'])
                    data['documents'].append(doc)
                
                with open(filename, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2, default=str)
            
            elif format == 'bibtex':
                # Export as BibTeX
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT d.id, d.filename 
                    FROM documents d
                """)
                
                with open(filename, 'w', encoding='utf-8') as f:
                    for doc_id, filename in cursor.fetchall():
                        meta = self.db.get_document_metadata(doc_id)
                        
                        # Generate BibTeX entry
                        entry_type = '@article'
                        if 'book' in filename.lower():
                            entry_type = '@book'
                        elif 'thesis' in filename.lower():
                            entry_type = '@phdthesis'
                        
                        cite_key = re.sub(r'[^\w]', '', filename)[:20]
                        
                        f.write(f"{entry_type}{{{cite_key},\n")
                        if meta.get('title'):
                            f.write(f"  title = {{{meta['title']}}},\n")
                        if meta.get('author'):
                            f.write(f"  author = {{{meta['author']}}},\n")
                        if meta.get('year'):
                            f.write(f"  year = {{{meta['year']}}},\n")
                        if meta.get('journal'):
                            f.write(f"  journal = {{{meta['journal']}}},\n")
                        f.write("}\n\n")
    
    def browse_output_dir(self):
        """Browse for output directory"""
        folder = filedialog.askdirectory()
        if folder:
            self.output_dir_var.set(folder)
    
    def save_settings(self):
        """Save application settings"""
        # Update config
        self.config['output_dir'] = self.output_dir_var.get()
        self.config['max_threads'] = self.max_threads_var.get()
        self.config['ocr_enabled'] = self.ocr_enabled_var.get()
        self.config['summary_method'] = self.summary_method_var.get()
        self.config['ollama_url'] = self.ollama_url_var.get()
        self.config['openai_api_key'] = self.openai_key_var.get()
        
        # Save to file
        try:
            with open('config.json', 'w') as f:
                json.dump(self.config, f, indent=2)
            
            messagebox.showinfo("Settings Saved", "Settings have been saved successfully")
        except Exception as e:
            messagebox.showerror("Save Error", f"Could not save settings: {str(e)}")
    
    def show_preferences(self):
        """Show preferences dialog"""
        # Switch to settings tab
        self.notebook.select(5)  # Settings is the 6th tab (0-indexed)
    
    def batch_process(self):
        """Batch process documents"""
        batch_window = tk.Toplevel(self.root)
        batch_window.title("Batch Processing")
        batch_window.geometry("600x400")
        
        ttk.Label(batch_window, text="Batch Document Processing", 
                 font=('Arial', 12, 'bold')).pack(pady=10)
        
        # Options
        options_frame = ttk.Frame(batch_window)
        options_frame.pack(fill='x', padx=20, pady=10)
        
        process_options = {
            'reprocess_all': tk.BooleanVar(value=False),
            'extract_citations': tk.BooleanVar(value=True),
            'generate_summaries': tk.BooleanVar(value=True),
            'update_metadata': tk.BooleanVar(value=True)
        }
        
        for i, (name, var) in enumerate(process_options.items()):
            ttk.Checkbutton(options_frame, text=name.replace('_', ' ').title(), 
                           variable=var).grid(row=i//2, column=i%2, sticky='w', padx=10, pady=5)
        
        # Progress
        progress_frame = ttk.Frame(batch_window)
        progress_frame.pack(fill='x', padx=20, pady=20)
        
        progress_label = ttk.Label(progress_frame, text="Ready to process")
        progress_label.pack()
        
        progress_bar = ttk.Progressbar(progress_frame, mode='determinate')
        progress_bar.pack(fill='x', pady=10)
        
        # Log
        log_frame = ttk.Frame(batch_window)
        log_frame.pack(fill='both', expand=True, padx=20, pady=10)
        
        log_text = scrolledtext.ScrolledText(log_frame, height=10)
        log_text.pack(fill='both', expand=True)
        
        def run_batch():
            # Get documents to process
            with sqlite3.connect(self.db.db_path) as conn:
                cursor = conn.cursor()
                
                if process_options['reprocess_all'].get():
                    cursor.execute("SELECT id, filepath FROM documents")
                else:
                    cursor.execute("SELECT id, filepath FROM documents WHERE processed = 0")
                
                docs = cursor.fetchall()
            
            if not docs:
                log_text.insert(tk.END, "No documents to process.\n")
                return
            
            progress_bar['maximum'] = len(docs)
            
            for i, (doc_id, filepath) in enumerate(docs):
                progress_label.config(text=f"Processing: {os.path.basename(filepath)}")
                log_text.insert(tk.END, f"Processing {filepath}...\n")
                log_text.see(tk.END)
                batch_window.update()
                
                try:
                    # Reprocess document
                    result = self.processor.process_document(filepath)
                    
                    if result.get('error'):
                        log_text.insert(tk.END, f"  Error: {result['error']}\n")
                    else:
                        log_text.insert(tk.END, "  Success\n")
                        
                        # Generate summary if requested
                        if process_options['generate_summaries'].get():
                            if result.get('full_text'):
                                summary = self.summarizer.summarize(
                                    result['full_text'][:5000], 
                                    length='medium'
                                )
                                # Store summary in database
                                with sqlite3.connect(self.db.db_path) as conn:
                                    cursor = conn.cursor()
                                    cursor.execute("""
                                        INSERT INTO summaries 
                                        (document_id, summary_type, summary_text, method_used)
                                        VALUES (?, ?, ?, ?)
                                    """, (doc_id, 'auto', summary, self.config['summary_method']))
                                    conn.commit()
                                
                                log_text.insert(tk.END, "  Generated summary\n")
                    
                except Exception as e:
                    log_text.insert(tk.END, f"  Error: {str(e)}\n")
                
                progress_bar['value'] = i + 1
                batch_window.update()
            
            progress_label.config(text="Batch processing complete")
            log_text.insert(tk.END, "\nBatch processing complete!\n")
            messagebox.showinfo("Complete", "Batch processing finished")
        
        # Buttons
        button_frame = ttk.Frame(batch_window)
        button_frame.pack(pady=10)
        
        ttk.Button(button_frame, text="Start Processing", 
                  command=run_batch, style='Accent.TButton').pack(side='left', padx=5)
        ttk.Button(button_frame, text="Close", 
                  command=batch_window.destroy).pack(side='left', padx=5)
    
    def citation_manager(self):
        """Open citation manager"""
        cite_window = tk.Toplevel(self.root)
        cite_window.title("Citation Manager")
        cite_window.geometry("800x600")
        
        ttk.Label(cite_window, text="Citation Manager", 
                 font=('Arial', 12, 'bold')).pack(pady=10)
        
        # Citation list
        columns = ('Document', 'Citation', 'Type', 'Year')
        cite_tree = ttk.Treeview(cite_window, columns=columns, show='headings')
        
        for col in columns:
            cite_tree.heading(col, text=col)
            cite_tree.column(col, width=150)
        
        cite_tree.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Load citations
        with sqlite3.connect(self.db.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT d.filename, c.citation_text, c.citation_type, c.year
                FROM citations c
                JOIN documents d ON c.document_id = d.id
                ORDER BY c.year DESC
            """)
            
            for row in cursor.fetchall():
                cite_tree.insert('', 'end', values=row)
        
        # Export button
        ttk.Button(cite_window, text="Export Citations", 
                  command=lambda: self.export_citations(cite_window)).pack(pady=10)
    
    def export_citations(self, parent_window):
        """Export citations"""
        filename = filedialog.asksaveasfilename(
            parent=parent_window,
            defaultextension=".bib",
            filetypes=[("BibTeX", "*.bib"), ("Text", "*.txt")]
        )
        
        if filename:
            with sqlite3.connect(self.db.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT citation_text, authors, year
                    FROM citations
                    ORDER BY year DESC
                """)
                
                with open(filename, 'w', encoding='utf-8') as f:
                    for citation, authors, year in cursor.fetchall():
                        f.write(f"{citation}\n")
            
            messagebox.showinfo("Export Complete", f"Citations exported to:\n{filename}")
    
    def start_web_server(self):
        """Start web interface server"""
        if not FLASK_AVAILABLE:
            messagebox.showerror("Flask Not Available", 
                               "Flask is not installed. Install with: pip install flask")
            return
        
        # Create and start Flask app in a thread
        def run_server():
            app = create_flask_app(self.config, self.db)
            app.run(host='127.0.0.1', port=5000, debug=False)
        
        server_thread = threading.Thread(target=run_server, daemon=True)
        server_thread.start()
        
        # Open browser
        time.sleep(1)
        webbrowser.open('http://127.0.0.1:5000')
        
        messagebox.showinfo("Web Server Started", 
                          "Web interface started at http://127.0.0.1:5000")
    
    def show_help(self):
        """Show help documentation"""
        help_window = tk.Toplevel(self.root)
        help_window.title("Help - Advanced Research System")
        help_window.geometry("800x600")
        
        help_text = scrolledtext.ScrolledText(help_window, wrap='word')
        help_text.pack(fill='both', expand=True, padx=10, pady=10)
        
        help_content = """ADVANCED RESEARCH SYSTEM - HELP

TABLE OF CONTENTS
1. Getting Started
2. Document Library
3. Search Functions
4. Analysis Tools
5. Web Research
6. Report Generation
7. Keyboard Shortcuts

1. GETTING STARTED
==================
The Advanced Research System (ARS) is a comprehensive tool for managing, searching, and analyzing research documents.

Quick Start:
- Click "Add Documents" or "Add Folder" to import your research files
- Use the Search tab to find specific content
- Run analyses to extract insights
- Generate professional reports

Supported File Types:
- PDF (.pdf)
- Word Documents (.docx)
- Text Files (.txt)
- Markdown (.md)
- HTML (.html, .htm)

2. DOCUMENT LIBRARY
==================
The Document Library is your central repository for all research materials.

Features:
- Import individual files or entire folders
- Automatic metadata extraction
- Full-text indexing
- Preview documents
- Filter by filename

Tips:
- Double-click a document to open it
- Use the filter box for quick searches
- Right-click for context menu options

3. SEARCH FUNCTIONS
==================
Three search modes are available:

Basic Search:
- Simple keyword matching
- Fast and straightforward

Advanced Search:
- Query operators: +must -exclude "exact phrase"
- Relevance ranking
- Metadata search

Semantic Search:
- Meaning-based matching
- Finds related concepts

Search Tips:
- Use quotes for exact phrases
- + prefix for required terms
- - prefix to exclude terms

4. ANALYSIS TOOLS
=================
Available analyses:

Statistics:
- Document counts
- Average lengths
- Type distribution

Theme Extraction:
- Common topics
- Keyword frequency
- Concept clustering

Citation Analysis:
- Extract citations
- Build bibliography
- Track references

5. WEB RESEARCH
===============
Extend your research with web sources:

Features:
- Web search integration
- Page crawling
- Content extraction
- Source tracking

Usage:
- Enter search terms
- Enable deep crawl for thorough results
- Preview results before saving

6. REPORT GENERATION
===================
Create professional research reports:

Formats:
- Markdown
- HTML
- PDF
- Word (DOCX)
- LaTeX

Sections:
- Executive Summary
- Introduction
- Methodology
- Literature Review
- Findings
- Analysis
- Conclusions
- References

7. KEYBOARD SHORTCUTS
====================
Ctrl+O - Open/Import documents
Ctrl+S - Save current work
Ctrl+F - Focus search box
Ctrl+R - Generate report
Ctrl+E - Export results
F5 - Refresh library
Esc - Close dialogs

For more information, visit the project documentation.
"""
        
        help_text.insert(1.0, help_content)
        help_text.config(state='disabled')
        
        ttk.Button(help_window, text="Close", 
                  command=help_window.destroy).pack(pady=10)
    
    def show_about(self):
        """Show about dialog"""
        about_text = f"""Advanced Research System (ARS)
Version {DEFAULT_CONFIG['version']}

A comprehensive research platform for:
• Literature review management
• Document analysis
• Web research
• Report generation

Created with Python and modern NLP technologies.

Libraries used:
• Document Processing: PyMuPDF, python-docx, pdfplumber
• Analysis: NLTK, pandas, BeautifulSoup4
• Interface: tkinter, Flask
• AI: Ollama, OpenAI (optional)

© 2024 Advanced Research System Team
"""
        
        messagebox.showinfo("About ARS", about_text)
    
    def run(self):
        """Start the GUI application"""
        self.root.mainloop()

def create_flask_app(config, db_manager):
    """Create Flask web application"""
    app = Flask(__name__)
    app.config['SECRET_KEY'] = 'research-system-secret-key'
    
    @app.route('/')
    def index():
        return '''
        <!DOCTYPE html>
        <html>
        <head>
            <title>Advanced Research System - Web Interface</title>
            <style>
                body {
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                    max-width: 1200px;
                    margin: 0 auto;
                    padding: 20px;
                    background: #f5f7fa;
                }
                .container {
                    background: white;
                    padding: 30px;
                    border-radius: 10px;
                    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                }
                h1 {
                    color: #2c3e50;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                }
                .section {
                    margin: 20px 0;
                    padding: 20px;
                    background: #f8f9fa;
                    border-radius: 5px;
                }
                .button {
                    display: inline-block;
                    padding: 10px 20px;
                    background: #3498db;
                    color: white;
                    text-decoration: none;
                    border-radius: 5px;
                    margin: 5px;
                }
                .button:hover {
                    background: #2980b9;
                }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>Advanced Research System - Web Interface</h1>
                
                <div class="section">
                    <h2>Document Library</h2>
                    <p>Access and manage your research documents.</p>
                    <a href="/documents" class="button">View Documents</a>
                    <a href="/upload" class="button">Upload Documents</a>
                </div>
                
                <div class="section">
                    <h2>Search</h2>
                    <p>Search across all your documents.</p>
                    <a href="/search" class="button">Search Documents</a>
                </div>
                
                <div class="section">
                    <h2>Analysis</h2>
                    <p>Analyze your research collection.</p>
                    <a href="/analysis" class="button">Run Analysis</a>
                </div>
                
                <div class="section">
                    <h2>Reports</h2>
                    <p>Generate comprehensive research reports.</p>
                    <a href="/reports" class="button">Create Report</a>
                </div>
            </div>
        </body>
        </html>
        '''
    
    @app.route('/documents')
    def documents():
        # Get documents from database
        with sqlite3.connect(db_manager.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT id, filename, filetype, filesize, date_added
                FROM documents
                ORDER BY date_added DESC
            """)
            docs = cursor.fetchall()
        
        html = '''
        <!DOCTYPE html>
        <html>
        <head>
            <title>Document Library</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                table { border-collapse: collapse; width: 100%; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #4CAF50; color: white; }
                tr:nth-child(even) { background-color: #f2f2f2; }
            </style>
        </head>
        <body>
            <h1>Document Library</h1>
            <table>
                <tr>
                    <th>ID</th>
                    <th>Filename</th>
                    <th>Type</th>
                    <th>Size</th>
                    <th>Date Added</th>
                </tr>
        '''
        
        for doc in docs:
            html += f'''
                <tr>
                    <td>{doc[0]}</td>
                    <td>{doc[1]}</td>
                    <td>{doc[2]}</td>
                    <td>{doc[3]}</td>
                    <td>{doc[4]}</td>
                </tr>
            '''
        
        html += '''
            </table>
            <p><a href="/">Back to Home</a></p>
        </body>
        </html>
        '''
        
        return html
    
    @app.route('/search', methods=['GET', 'POST'])
    def search():
        if request.method == 'POST':
            query = request.form.get('query', '')
            # Perform search using the search engine
            # Return results
            return f"<h1>Search Results for: {query}</h1><p>Search functionality to be implemented</p>"
        
        return '''
        <!DOCTYPE html>
        <html>
        <head>
            <title>Search Documents</title>
        </head>
        <body>
            <h1>Search Documents</h1>
            <form method="post">
                <input type="text" name="query" placeholder="Enter search query" size="50">
                <button type="submit">Search</button>
            </form>
            <p><a href="/">Back to Home</a></p>
        </body>
        </html>
        '''
    
    return app

def main():
    """Main entry point for the application"""
    # Create necessary directories
    for dir_name in ['research_output', 'temp', 'cache']:
        os.makedirs(dir_name, exist_ok=True)
    
    parser = argparse.ArgumentParser(
        description="Advanced Research System - Ultimate Literature Review & Analysis Tool",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Launch GUI
  python ars.py

  # Process documents from command line
  python ars.py --process-folder /path/to/documents --output-dir ./results

  # Run web search
  python ars.py --web-search "machine learning healthcare" --max-results 50

  # Generate report
  python ars.py --generate-report --format pdf --output report.pdf

  # Start web server
  python ars.py --web-server --port 5000
        """
    )
    
    # Mode selection
    parser.add_argument('--gui', action='store_true', default=True,
                       help='Launch GUI interface (default)')
    parser.add_argument('--cli', action='store_true',
                       help='Use command-line interface')
    parser.add_argument('--web-server', action='store_true',
                       help='Start web server interface')
    
    # Processing options
    parser.add_argument('--process-file', type=str,
                       help='Process a single document')
    parser.add_argument('--process-folder', type=str,
                       help='Process all documents in a folder')
    parser.add_argument('--recursive', action='store_true',
                       help='Process folders recursively')
    
    # Search options
    parser.add_argument('--search', type=str,
                       help='Search query')
    parser.add_argument('--search-type', choices=['basic', 'advanced', 'semantic'],
                       default='advanced', help='Search type')
    parser.add_argument('--web-search', type=str,
                       help='Web search query')
    parser.add_argument('--max-results', type=int, default=30,
                       help='Maximum search results')
    
    # Analysis options
    parser.add_argument('--analyze', action='store_true',
                       help='Run analysis on documents')
    parser.add_argument('--extract-themes', action='store_true',
                       help='Extract themes from documents')
    parser.add_argument('--citation-analysis', action='store_true',
                       help='Analyze citations')
    
    # Report options
    parser.add_argument('--generate-report', action='store_true',
                       help='Generate research report')
    parser.add_argument('--report-title', type=str, default='Research Report',
                       help='Report title')
    parser.add_argument('--format', choices=['markdown', 'html', 'pdf', 'docx', 'latex'],
                       default='markdown', help='Output format')
    
    # Output options
    parser.add_argument('--output', type=str,
                       help='Output file path')
    parser.add_argument('--output-dir', type=str, default='research_output',
                       help='Output directory')
    
    # Configuration
    parser.add_argument('--config', type=str,
                       help='Configuration file path')
    parser.add_argument('--db-path', type=str, default='research_database.db',
                       help='Database file path')
    
    # Server options
    parser.add_argument('--port', type=int, default=5000,
                       help='Web server port')
    parser.add_argument('--host', type=str, default='127.0.0.1',
                       help='Web server host')
    
    args = parser.parse_args()
    
    # Load configuration
    config = DEFAULT_CONFIG.copy()
    if args.config and os.path.exists(args.config):
        with open(args.config, 'r') as f:
            config.update(json.load(f))
    
    # Update config with command line arguments
    config['database_path'] = args.db_path
    config['output_dir'] = args.output_dir
    
    # Create output directory
    os.makedirs(config['output_dir'], exist_ok=True)
    
    # Initialize database
    db_manager = DatabaseManager(config['database_path'])
    
    # Check if any CLI operations requested
    cli_mode = any([
        args.cli, args.process_file, args.process_folder,
        args.search, args.web_search, args.analyze,
        args.generate_report, args.web_server
    ])
    
    if args.web_server:
        # Start web server
        if not FLASK_AVAILABLE:
            print("Error: Flask not installed. Install with: pip install flask")
            sys.exit(1)
        
        app = create_flask_app(config, db_manager)
        print(f"Starting web server at http://{args.host}:{args.port}")
        app.run(host=args.host, port=args.port, debug=False)
    
    elif cli_mode and not args.gui:
        # CLI mode
        print("Advanced Research System - CLI Mode")
        print("="*50)
        
        # Initialize components
        processor = DocumentProcessor(config, db_manager)
        search_engine = AdvancedSearchEngine(config, db_manager)
        web_engine = WebResearchEngine(config, db_manager)
        summarizer = SummarizationEngine(config)
        report_gen = ReportGenerator(config)
        
        # Process files
        if args.process_file:
            print(f"Processing file: {args.process_file}")
            result = processor.process_document(args.process_file)
            if result.get('error'):
                print(f"Error: {result['error']}")
            else:
                print(f"Successfully processed: {args.process_file}")
        
        elif args.process_folder:
            print(f"Processing folder: {args.process_folder}")
            
            # Find all supported files
            files = []
            if args.recursive:
                for ext in config['supported_formats']:
                    files.extend(Path(args.process_folder).rglob(f"*{ext}"))
            else:
                for ext in config['supported_formats']:
                    files.extend(Path(args.process_folder).glob(f"*{ext}"))
            
            print(f"Found {len(files)} files to process")
            
            # Process each file
            for i, file in enumerate(files, 1):
                print(f"Processing ({i}/{len(files)}): {file.name}")
                try:
                    result = processor.process_document(str(file))
                    if result.get('error'):
                        print(f"  Error: {result['error']}")
                except Exception as e:
                    print(f"  Error: {str(e)}")
            
            print("Processing complete")
        
        # Search
        if args.search:
            print(f"Searching for: {args.search}")
            results = search_engine.search(args.search, args.search_type)
            
            print(f"Found {len(results)} results:")
            for i, result in enumerate(results[:10], 1):
                print(f"\n{i}. {result['filename']} (Score: {result['relevance_score']:.2f})")
                if result.get('paragraph_content'):
                    preview = result['paragraph_content'][:200] + "..."
                    print(f"   {preview}")
        
        # Web search
        if args.web_search:
            print(f"Searching web for: {args.web_search}")
            results = web_engine.search_web(args.web_search, args.max_results)
            
            print(f"Found {len(results)} web results:")
            for i, result in enumerate(results[:10], 1):
                print(f"\n{i}. {result['title']}")
                print(f"   URL: {result['url']}")
                print(f"   {result['snippet']}")
        
        # Analysis
        if args.analyze or args.extract_themes or args.citation_analysis:
            print("Running analysis...")
            
            # Get document count
            with sqlite3.connect(db_manager.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT COUNT(*) FROM documents WHERE processed = 1")
                doc_count = cursor.fetchone()[0]
            
            print(f"Analyzing {doc_count} documents")
            
            # Run requested analyses
            # (Implementation would go here)
            print("Analysis complete")
        
        # Generate report
        if args.generate_report:
            print("Generating report...")
            
            # Gather report data
            report_data = {
                'title': args.report_title,
                'documents': []
                # (Additional data gathering would go here)
            }
            
            # Generate report
            report_content = report_gen.generate_report(report_data, args.format)
            
            # Save report
            output_file = args.output or f"report.{args.format}"
            
            if args.format in ['markdown', 'html', 'latex']:
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(report_content)
            else:
                with open(output_file, 'wb') as f:
                    f.write(report_content)
            
            print(f"Report saved to: {output_file}")
    
    else:
        # GUI mode (default)
        if not GUI_AVAILABLE:
            print("Error: Tkinter not available. Use --cli flag for command-line mode.")
            sys.exit(1)
        
        print("Launching Advanced Research System GUI...")
        app = AdvancedResearchGUI()
        app.run()

if __name__ == "__main__":
    main()